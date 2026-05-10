import tkinter as tk
import tkinter.messagebox as messagebox
from PIL import Image, ImageTk
import os
import sys
import sqlite3
import uuid
import tkinter.filedialog as filedialog
import json
import shutil
from datetime import datetime
import subprocess
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import calendar
from datetime import date


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


APP_DATA_DIR = os.path.join(os.path.expanduser("~"), "IDARA_DZ")
CLIENTS_DATA_FILE = os.path.join(APP_DATA_DIR, "clients_data.json")
ARCHIVE_DIR = os.path.join(APP_DATA_DIR, "archive")

os.makedirs(APP_DATA_DIR, exist_ok=True)
os.makedirs(ARCHIVE_DIR, exist_ok=True)


def load_clients_data():
    if not os.path.exists(CLIENTS_DATA_FILE):
        return []

    try:
        with open(CLIENTS_DATA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except:
        return []


def save_clients_data(data):
    with open(CLIENTS_DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


def save_or_update_client_from_form(generated_file_path=None):
    clients = load_clients_data()

    first_name = job_form_entries.get("first_name", "").strip()
    last_name = job_form_entries.get("last_name", "").strip()
    birth_info = job_form_entries.get("birth_info", "").strip()
    birth_info = job_form_entries.get("birth_info", "").strip()
    address = job_form_entries.get("address", "").strip()
    phone = job_form_entries.get("phone", "").strip()
    id_card = job_form_entries.get("id_card", "").strip()

    if not first_name and not last_name and not phone and not id_card:
        return

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    request_type = job_form_entries.get("request_type", "").strip() or "طلب توظيف (عام)"

    found = None

    for client in clients:
        same_phone = phone and client.get("phone") == phone
        same_id = id_card and client.get("id_card") == id_card
        same_name = first_name and last_name and client.get("first_name") == first_name and client.get("last_name") == last_name

        if same_phone or same_id or same_name:
            found = client
            break

    record = {
        "first_name": first_name,
        "last_name": last_name,
        "birth_info": birth_info,
        "address": address,
        "phone": phone,
        "id_card": id_card,
        "last_request": request_type,
        "last_used": now,
        "last_file": generated_file_path or ""
    }

    if found:
        found.update(record)
    else:
        clients.append(record)

    save_clients_data(clients)


def archive_generated_file(file_path):
    if not file_path or not os.path.exists(file_path):
        return file_path

    request_type = job_form_entries.get("request_type", "").strip() or "طلب توظيف (عام)"
    first_name = job_form_entries.get("first_name", "").strip()
    last_name = job_form_entries.get("last_name", "").strip()
    full_name = f"{last_name}_{first_name}".strip("_") or "بدون_اسم"

    safe_request = request_type.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    safe_name = full_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")

    folder = os.path.join(ARCHIVE_DIR, "طلبات خطية", safe_request)
    os.makedirs(folder, exist_ok=True)

    stamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    archive_path = os.path.join(folder, f"{safe_name}_{stamp}.docx")

    shutil.copy2(file_path, archive_path)
    return archive_path


def find_client_suggestions(text):
    query = (text or "").strip().lower()
    if not query:
        return []

    results = []
    for client in load_clients_data():
        combined = " ".join([
            client.get("first_name", ""),
            client.get("last_name", ""),
            client.get("phone", ""),
            client.get("id_card", ""),
        ]).lower()

        if query in combined:
            results.append(client)

    return results[:5]


def apply_client_to_form(client):
    job_form_entries["first_name"] = client.get("first_name", "")
    job_form_entries["last_name"] = client.get("last_name", "")
    job_form_entries["birth_info"] = client.get("birth_info", "")
    job_form_entries["address"] = client.get("address", "")
    job_form_entries["phone"] = client.get("phone", "")
    job_form_entries["id_card"] = client.get("id_card", "")
    show_job_request_form()



root = tk.Tk()
root.title("IDARA DZ")
root.geometry("1280x720")
root.minsize(1000, 600)
root.configure(bg="black")

canvas = tk.Canvas(root, highlightthickness=0, bd=0)
canvas.pack(fill="both", expand=True)

background_photo = None
home_card_photos = {}
document_card_photos = {}
current_page = "home"
animation_running = False


def clear_screen():
    global animation_running
    animation_running = False
    canvas.delete("all")


def draw_image(path, width, height):
    global background_photo

    if os.path.exists(path):
        img = Image.open(path).convert("RGB")
        img = img.resize((width, height), Image.LANCZOS)
    else:
        img = Image.new("RGB", (width, height), "#173b38")

    background_photo = ImageTk.PhotoImage(img)
    canvas.create_image(0, 0, image=background_photo, anchor="nw")


def draw_home_sidebar(active="home"):
    width = root.winfo_width()
    height = root.winfo_height()
    sidebar_w = 120

    canvas.create_rectangle(0, 0, sidebar_w, height, fill="#f7f7f7", outline="#e5e5e5")
    canvas.create_line(sidebar_w, 0, sidebar_w, height, fill="#dedede", width=1)

    canvas.create_rectangle(20, 22, 42, 44, fill="#000000", outline="#000000")
    canvas.create_text(52, 33, text="IDARA DZ", fill="#111111", font=("Arial", 12, "bold"), anchor="w")

    items = [
        ("⌂", "الرئيسية", "home", 105),
        ("⚙", "الإعدادات", "settings", 205),
        ("ⓘ", "حول البرنامج", "about", 305),
    ]

    for icon, label, key, y in items:
        is_active = active == key

        if is_active:
            bg = canvas.create_rectangle(8, y - 48, sidebar_w - 8, y + 44, fill="#000000", outline="#000000")
            fill = "#ffffff"
        else:
            bg = canvas.create_rectangle(8, y - 48, sidebar_w - 8, y + 44, fill="#f7f7f7", outline="#f7f7f7")
            fill = "#000000"

        icon_id = canvas.create_text(sidebar_w // 2, y - 14, text=icon, fill=fill, font=("Arial", 28, "bold"))
        label_id = canvas.create_text(sidebar_w // 2, y + 20, text=label, fill=fill, font=("Arial", 11, "bold"))

        def enter(event, b=bg, active_state=is_active):
            if not active_state:
                canvas.itemconfig(b, fill="#ededed", outline="#ededed")
            root.config(cursor="hand2")

        def leave(event, b=bg, active_state=is_active):
            if not active_state:
                canvas.itemconfig(b, fill="#f7f7f7", outline="#f7f7f7")
            root.config(cursor="")

        def click(event, target=key):
            if target == "home":
                show_home()
            elif target == "about":
                show_about()
            else:
                show_settings_main()

        for item in (bg, icon_id, label_id):
            canvas.tag_bind(item, "<Enter>", enter)
            canvas.tag_bind(item, "<Leave>", leave)
            canvas.tag_bind(item, "<Button-1>", click)

    dark_y = height - 145
    canvas.create_text(sidebar_w // 2, dark_y, text="☾", fill="#000000", font=("Arial", 26, "bold"))
    canvas.create_text(sidebar_w // 2, dark_y + 32, text="الوضع الداكن", fill="#000000", font=("Arial", 10, "bold"))
    canvas.create_line(16, height - 102, sidebar_w - 16, height - 102, fill="#dddddd")

    exit_y = height - 62
    exit_icon = canvas.create_text(sidebar_w // 2, exit_y - 12, text="↪", fill="#000000", font=("Arial", 24, "bold"))
    exit_label = canvas.create_text(sidebar_w // 2, exit_y + 18, text="خروج", fill="#000000", font=("Arial", 10, "bold"))

    def exit_enter(event):
        root.config(cursor="hand2")

    def exit_leave(event):
        root.config(cursor="")

    def exit_click(event):
        root.destroy()

    for item in (exit_icon, exit_label):
        canvas.tag_bind(item, "<Enter>", exit_enter)
        canvas.tag_bind(item, "<Leave>", exit_leave)
        canvas.tag_bind(item, "<Button-1>", exit_click)


def rounded_home_rect(x1, y1, x2, y2, r=16, fill="#ffffff", outline="#dddddd", width=1):
    points = [
        x1 + r, y1, x2 - r, y1, x2, y1, x2, y1 + r,
        x2, y2 - r, x2, y2, x2 - r, y2, x1 + r, y2,
        x1, y2, x1, y2 - r, x1, y1 + r, x1, y1
    ]
    return canvas.create_polygon(points, smooth=True, fill=fill, outline=outline, width=width)


def show_home():
    global current_page
    current_page = "home"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    sidebar_w = 120
    content_x1 = sidebar_w
    content_w = width - sidebar_w
    center_x = content_x1 + content_w // 2

    logo_y = int(height * 0.22)

    # Logo image from assets/logo.png
    try:
        logo_img = Image.open(resource_path("assets/logo.png")).convert("RGBA")
        logo_img = logo_img.resize((140, 140), Image.LANCZOS)
        logo_photo = ImageTk.PhotoImage(logo_img)
        canvas.logo_photo = logo_photo
        canvas.create_image(center_x - 165, logo_y, image=logo_photo)
    except Exception:
        canvas.create_rectangle(center_x - 245, logo_y - 54, center_x - 160, logo_y + 54, fill="#000000", outline="#000000")
        canvas.create_text(center_x - 202, logo_y, text="▰", fill="#ffffff", font=("Arial", 38, "bold"))

    canvas.create_text(center_x - 105, logo_y, text="IDARA", fill="#000000", font=("Arial", 48, "bold"), anchor="w")
    canvas.create_text(center_x + 95, logo_y, text="DZ", fill="#777777", font=("Arial", 48, "bold"), anchor="w")

    subtitle_y = logo_y + 85
    canvas.create_line(center_x - 190, subtitle_y, center_x - 120, subtitle_y, fill="#bbbbbb")
    canvas.create_text(center_x, subtitle_y, text="خدمات إدارية بكل احترافية", fill="#555555", font=("Arial", 18, "bold"))
    canvas.create_line(center_x + 120, subtitle_y, center_x + 190, subtitle_y, fill="#bbbbbb")

    cards = [
        ("assets/documents.png", "وثائق", "إنشاء وتعديل مختلف الوثائق\nالإدارية بسهولة", "documents"),
        ("assets/electronic.png", "خدمات الكترونية", "الوصول إلى الخدمات الإلكترونية\nوالمنصات الرسمية", "electronic"),
        ("assets/archive.png", "ارشيف", "إدارة وأرشفة الملفات والوثائق\nوالوصول إليها بسهولة", "archive"),
    ]

    card_w = int(content_w * 0.245)
    card_h = int(height * 0.39)
    gap = int(content_w * 0.04)
    total_w = card_w * 3 + gap * 2
    start_x = center_x - total_w // 2
    card_y1 = int(height * 0.40)
    card_y2 = card_y1 + card_h

    global home_card_photos
    home_card_photos = {}

    for i, (icon_path, title, desc, key) in enumerate(cards):
        x1 = start_x + i * (card_w + gap)
        x2 = x1 + card_w

        card = rounded_home_rect(x1, card_y1, x2, card_y2, r=14, fill="#ffffff", outline="#dddddd", width=1)

        try:
            icon_img = Image.open(resource_path(icon_path)).convert("RGBA")
            icon_img = icon_img.resize((95, 95), Image.LANCZOS)
            icon_photo = ImageTk.PhotoImage(icon_img)
            home_card_photos[key] = icon_photo
            icon_id = canvas.create_image((x1 + x2) // 2, card_y1 + 82, image=icon_photo)
        except Exception:
            icon_id = canvas.create_text((x1 + x2) // 2, card_y1 + 82, text="■", fill="#000000", font=("Arial", 54, "bold"))

        title_id = canvas.create_text((x1 + x2) // 2, card_y1 + 165, text=title, fill="#000000", font=("Arial", 27, "bold"))
        desc_id = canvas.create_text((x1 + x2) // 2, card_y1 + 225, text=desc, fill="#666666", font=("Arial", 15, "bold"), justify="center")

        def enter(event, c=card):
            canvas.itemconfig(c, fill="#fafafa", outline="#cfcfcf")
            root.config(cursor="hand2")

        def leave(event, c=card):
            canvas.itemconfig(c, fill="#ffffff", outline="#dddddd")
            root.config(cursor="")

        def click(event, k=key):
            if k == "documents":
                show_documents()
            else:
                show_section(k)

        for item in (card, icon_id, title_id, desc_id):
            canvas.tag_bind(item, "<Enter>", enter)
            canvas.tag_bind(item, "<Leave>", leave)
            canvas.tag_bind(item, "<Button-1>", click)

    canvas.create_text(center_x, height - 55, text="© 2024 IDARA DZ - جميع الحقوق محفوظة", fill="#777777", font=("Arial", 13))


def show_settings_main():
    show_settings_main()


def draw_back_button(command):
    btn = canvas.create_rectangle(
        35,
        35,
        145,
        82,
        fill="#173b38",
        outline="#d7c28a",
        width=2
    )

    txt = canvas.create_text(
        90,
        58,
        text="رجوع",
        fill="#f4f4f4",
        font=("Arial", 16, "bold")
    )

    def enter(event):
        canvas.itemconfig(btn, fill="#21524d")
        root.config(cursor="hand2")

    def leave(event):
        canvas.itemconfig(btn, fill="#173b38")
        root.config(cursor="")

    def click(event):
        command()

    for item in (btn, txt):
        canvas.tag_bind(item, "<Enter>", enter)
        canvas.tag_bind(item, "<Leave>", leave)
        canvas.tag_bind(item, "<Button-1>", click)


def show_documents():
    global current_page, document_card_photos
    current_page = "documents"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    sidebar_w = 120
    content_x1 = sidebar_w
    content_w = width - sidebar_w
    center_x = content_x1 + content_w // 2

    header_y = int(height * 0.16)

    try:
        icon_img = Image.open(resource_path("assets/documents.png")).convert("RGBA")
        icon_img = icon_img.resize((92, 92), Image.LANCZOS)
        doc_header_photo = ImageTk.PhotoImage(icon_img)
        canvas.doc_header_photo = doc_header_photo
        canvas.create_image(center_x - 70, header_y, image=doc_header_photo)
    except Exception:
        canvas.create_text(center_x - 70, header_y, text="▣", fill="#000000", font=("Arial", 64, "bold"))

    canvas.create_text(
        center_x + 40,
        header_y - 3,
        text="وثائق",
        fill="#000000",
        font=("Arial", 54, "bold")
    )

    canvas.create_text(
        center_x,
        header_y + 58,
        text="— إنشاء و تعديل مختلف الوثائق الإدارية —",
        fill="#111111",
        font=("Arial", 15, "bold")
    )

    cards = [
        ("assets/written_request.png", "طلب خطي", "written_request"),
        ("assets/honor_statement.png", "تصريح شرفي", "honor_statement"),
        ("assets/cv.png", "سيرة ذاتية", "cv"),
        ("assets/invoice.png", "فاتورة", "invoice"),
    ]

    document_card_photos = {}

    card_w = int(content_w * 0.19)
    card_h = int(height * 0.34)
    gap = int(content_w * 0.055)
    total_w = card_w * 4 + gap * 3
    start_x = center_x - total_w // 2

    card_y1 = int(height * 0.37)
    card_y2 = card_y1 + card_h

    for i, (icon_path, title, key) in enumerate(cards):
        x1 = start_x + i * (card_w + gap)
        x2 = x1 + card_w

        shadow = rounded_home_rect(
            x1 + 8,
            card_y1 + 12,
            x2 + 8,
            card_y2 + 12,
            r=10,
            fill="#d9d9d9",
            outline="#d9d9d9",
            width=1
        )

        card = rounded_home_rect(
            x1,
            card_y1,
            x2,
            card_y2,
            r=10,
            fill="#ffffff",
            outline="#eeeeee",
            width=1
        )

        try:
            item_img = Image.open(resource_path(icon_path)).convert("RGBA")
            item_img = item_img.resize((115, 115), Image.LANCZOS)
            item_photo = ImageTk.PhotoImage(item_img)
            document_card_photos[key] = item_photo
            icon_id = canvas.create_image((x1 + x2) // 2, card_y1 + 88, image=item_photo)
        except Exception:
            icon_id = canvas.create_text(
                (x1 + x2) // 2,
                card_y1 + 88,
                text="▣",
                fill="#000000",
                font=("Arial", 58, "bold")
            )

        title_id = canvas.create_text(
            (x1 + x2) // 2,
            card_y1 + 195,
            text=title,
            fill="#000000",
            font=("Arial", 31, "bold")
        )

        hitbox = canvas.create_rectangle(
            x1,
            card_y1,
            x2,
            card_y2,
            fill="",
            outline=""
        )

        def enter(event, c=card):
            canvas.itemconfig(c, fill="#fafafa", outline="#dddddd")
            root.config(cursor="hand2")

        def leave(event, c=card):
            canvas.itemconfig(c, fill="#ffffff", outline="#eeeeee")
            root.config(cursor="")

        def click(event, k=key):
            if k == "written_request":
                show_written_request_menu()
            else:
                show_document_type(k)

        for item in (card, icon_id, title_id, hitbox):
            canvas.tag_bind(item, "<Enter>", enter)
            canvas.tag_bind(item, "<Leave>", leave)
            canvas.tag_bind(item, "<Button-1>", click)


def show_document_type(doc_type):
    if doc_type == "written_request":
        show_written_request()
        return

    global current_page
    current_page = doc_type
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    draw_image(resource_path("assets/background.jpg"), width, height)

    titles = {
        "honor_statement": "تصريح شرفي",
        "cv": "سيرة ذاتية",
        "invoice": "فاتورة",
    }

    canvas.create_rectangle(
        int(width * 0.18),
        int(height * 0.22),
        int(width * 0.82),
        int(height * 0.72),
        fill="#173b38",
        outline="#d7c28a",
        width=2
    )

    canvas.create_text(
        width // 2,
        int(height * 0.34),
        text=titles.get(doc_type, "وثيقة"),
        fill="#f4f4f4",
        font=("Arial", 44, "bold")
    )

    canvas.create_text(
        width // 2,
        int(height * 0.50),
        text="سنقوم ببناء هذه الوثيقة لاحقًا.",
        fill="#e8e1d5",
        font=("Arial", 24, "bold")
    )

    draw_back_button(show_documents)


written_request_scroll = 0
written_request_items = [
    "طلب توظيف عام"
]


def show_written_request():
    global current_page, written_request_scroll
    current_page = "written_request"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    draw_image(resource_path("assets/background.jpg"), width, height)

    right_x = int(width * 0.58)
    menu_center_x = right_x + (width - right_x) // 2

    canvas.create_text(
        menu_center_x,
        int(height * 0.12),
        text="طلب خطي",
        fill="#f4f4f4",
        font=("Arial", 50, "bold"),
        anchor="center"
    )

    visible_count = 10
    max_scroll = max(0, len(written_request_items) - visible_count)
    written_request_scroll = max(0, min(written_request_scroll, max_scroll))

    start_y = int(height * 0.24)
    row_gap = int(height * 0.065)

    visible_items = written_request_items[written_request_scroll:written_request_scroll + visible_count]

    for index, label in enumerate(visible_items):
        y = start_y + index * row_gap

        line_id = canvas.create_line(
            right_x + 70,
            y - 18,
            width - 70,
            y - 18,
            fill="#8aa09c",
            width=1
        )

        text_id = canvas.create_text(
            menu_center_x,
            y,
            text=label,
            fill="#e8e1d5",
            font=("Arial", 23, "bold"),
            anchor="center"
        )

        hitbox = canvas.create_rectangle(
            right_x + 55,
            y - 25,
            width - 55,
            y + 25,
            fill="",
            outline=""
        )

        def on_enter(event, t=text_id):
            canvas.itemconfig(t, fill="#d7c28a")
            root.config(cursor="hand2")

        def on_leave(event, t=text_id):
            canvas.itemconfig(t, fill="#e8e1d5")
            root.config(cursor="")

        def on_click(event, name=label):
            show_written_request_template(name)

        for item in (text_id, hitbox, line_id):
            canvas.tag_bind(item, "<Enter>", on_enter)
            canvas.tag_bind(item, "<Leave>", on_leave)
            canvas.tag_bind(item, "<Button-1>", on_click)

    if written_request_scroll < max_scroll:
        canvas.create_text(
            menu_center_x,
            height - 42,
            text="∨",
            fill="#777777",
            font=("Arial", 34, "bold")
        )

    if written_request_scroll > 0:
        canvas.create_text(
            menu_center_x,
            int(height * 0.18),
            text="∧",
            fill="#777777",
            font=("Arial", 23, "bold")
        )

    draw_back_button(show_documents)


def scroll_written_request(event):
    global written_request_scroll

    if current_page != "written_request":
        return

    if event.delta < 0:
        written_request_scroll += 1
    else:
        written_request_scroll -= 1

    max_scroll = max(0, len(written_request_items) - 10)
    written_request_scroll = max(0, min(written_request_scroll, max_scroll))
    show_written_request()


def show_written_request_template(name):
    if name == "طلب توظيف عام":
        show_job_request_form()
        return

    global current_page
    current_page = "written_request_template"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    draw_image(resource_path("assets/background.jpg"), width, height)

    canvas.create_rectangle(
        int(width * 0.18),
        int(height * 0.22),
        int(width * 0.82),
        int(height * 0.72),
        fill="#173b38",
        outline="#d7c28a",
        width=2
    )

    canvas.create_text(
        width // 2,
        int(height * 0.34),
        text=name,
        fill="#f4f4f4",
        font=("Arial", 38, "bold")
    )

    canvas.create_text(
        width // 2,
        int(height * 0.50),
        text="سنقوم ببناء نموذج هذا الطلب في الخطوة القادمة.",
        fill="#e8e1d5",
        font=("Arial", 22, "bold")
    )

    draw_back_button(show_written_request)


job_form_scroll = 0
job_form_entries = {}

REQUEST_TYPE_OPTIONS = [
    "مسابقة على أساس الشهادة",
    "مسابقة على أساس الاختبار",
    "طلب استخلاف",
    "مسابقة الحماية المدنية",
    "مسابقة الشرطة",
    "مسابقة الجمارك",
    "طلب توظيف (عام)",
    "طلب استقالة",
    "طلب سكن",
    "طلب شهادة عمل",
    "طلب تربص ميداني",
    "طلب عطلة سنوية",
    "طلب تحويل إداري",
    "طلب تسوية وضعية",
]

request_type_dropdown_open = False
request_type_scroll = 0
request_type_selected_index = 6


def rounded_rect(x1, y1, x2, y2, r=24, fill="#f3eeee", outline="#f3eeee", width=1):
    points = [
        x1 + r, y1,
        x2 - r, y1,
        x2, y1,
        x2, y1 + r,
        x2, y2 - r,
        x2, y2,
        x2 - r, y2,
        x1 + r, y2,
        x1, y2,
        x1, y2 - r,
        x1, y1 + r,
        x1, y1
    ]
    return canvas.create_polygon(points, smooth=True, fill=fill, outline=outline, width=width)


suggestion_items = []


def clear_suggestions():
    global suggestion_items
    for item in suggestion_items:
        try:
            canvas.delete(item)
        except:
            pass
    suggestion_items = []


def draw_client_suggestions(x, y, w, field_key, current_text):
    global suggestion_items
    clear_suggestions()

    if field_key not in ["first_name", "last_name", "phone", "id_card"]:
        return

    results = find_client_suggestions(current_text)
    if not results:
        return

    row_h = 38
    box_h = row_h * len(results) + 10
    box_x1 = x - w // 2
    box_x2 = x + w // 2
    box_y1 = y + 36
    box_y2 = box_y1 + box_h

    bg = rounded_rect(
        box_x1,
        box_y1,
        box_x2,
        box_y2,
        r=18,
        fill="#f3eeee",
        outline="#d9d1d1",
        width=2
    )
    suggestion_items.append(bg)

    for index, client in enumerate(results):
        item_y = box_y1 + 10 + index * row_h + row_h // 2
        label = f"{client.get('last_name', '')} {client.get('first_name', '')} - {client.get('phone', '')}".strip()

        text = canvas.create_text(
            box_x2 - 20,
            item_y,
            text=label,
            fill="#173b38",
            font=("Arial", 14, "bold"),
            anchor="e"
        )

        hitbox = canvas.create_rectangle(
            box_x1 + 5,
            item_y - 17,
            box_x2 - 5,
            item_y + 17,
            fill="",
            outline=""
        )

        suggestion_items.extend([text, hitbox])

        def enter(event):
            root.config(cursor="hand2")

        def leave(event):
            root.config(cursor="")

        def click(event, selected=client):
            clear_suggestions()
            apply_client_to_form(selected)

        for item in (text, hitbox):
            canvas.tag_bind(item, "<Enter>", enter)
            canvas.tag_bind(item, "<Leave>", leave)
            canvas.tag_bind(item, "<Button-1>", click)


def make_placeholder_entry(x, y, w, h, placeholder, field_key):
    value = job_form_entries.get(field_key, "")

    rounded_rect(
        x - w // 2,
        y - h // 2,
        x + w // 2,
        y + h // 2,
        r=28,
        fill="#f3eeee",
        outline="#d9d1d1",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 21, "bold"),
        justify="right",
        bd=0,
        bg="#f3eeee",
        fg="#111111",
        insertbackground="#173b38"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#777777")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#777777")
            job_form_entries[field_key] = ""
            root.after(200, clear_suggestions)

    def save_value(event=None):
        val = entry.get()
        if val == placeholder:
            job_form_entries[field_key] = ""
            clear_suggestions()
        else:
            job_form_entries[field_key] = val
            draw_client_suggestions(x, y, w, field_key, val)

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w - 46, height=h - 16)
    return entry


def open_request_type_dropdown():
    global request_type_dropdown_open
    request_type_dropdown_open = True
    show_job_request_form()


def close_request_type_dropdown():
    global request_type_dropdown_open
    request_type_dropdown_open = False
    show_job_request_form()


def choose_request_type(value):
    global request_type_dropdown_open
    job_form_entries["request_type"] = value
    request_type_dropdown_open = False
    show_job_request_form()


def scroll_request_type_dropdown(direction):
    global request_type_scroll

    max_visible = 6
    max_scroll = max(0, len(REQUEST_TYPE_OPTIONS) - max_visible)

    if direction > 0:
        request_type_scroll += 1
    else:
        request_type_scroll -= 1

    request_type_scroll = max(0, min(request_type_scroll, max_scroll))
    show_job_request_form()


def move_request_type_selection(direction):
    global request_type_selected_index, request_type_scroll, request_type_dropdown_open

    if not request_type_dropdown_open:
        return

    request_type_selected_index += direction
    request_type_selected_index = max(0, min(request_type_selected_index, len(REQUEST_TYPE_OPTIONS) - 1))

    max_visible = 6
    if request_type_selected_index < request_type_scroll:
        request_type_scroll = request_type_selected_index
    elif request_type_selected_index >= request_type_scroll + max_visible:
        request_type_scroll = request_type_selected_index - max_visible + 1

    show_job_request_form()


def confirm_request_type_selection():
    global request_type_dropdown_open

    if not request_type_dropdown_open:
        return

    choose_request_type(REQUEST_TYPE_OPTIONS[request_type_selected_index])


def draw_request_type_field(x, y, w, h):
    selected_value = job_form_entries.get("request_type", "")

    rounded_rect(
        x - w // 2,
        y - h // 2,
        x + w // 2,
        y + h // 2,
        r=28,
        fill="#f3eeee",
        outline="#d9d1d1",
        width=2
    )

    arrow_x = x - w // 2 + 45
    arrow = canvas.create_text(
        arrow_x,
        y,
        text="▼",
        fill="#555555",
        font=("Arial", 26, "bold")
    )

    text = canvas.create_text(
        x + w // 2 - 48,
        y,
        text=selected_value if selected_value else "نوع الطلب",
        fill="#173b38" if selected_value else "#777777",
        font=("Arial", 21, "bold"),
        anchor="e"
    )

    hitbox = canvas.create_rectangle(
        x - w // 2,
        y - h // 2,
        x + w // 2,
        y + h // 2,
        fill="",
        outline=""
    )

    def enter(event):
        root.config(cursor="hand2")

    def leave(event):
        root.config(cursor="")

    def click(event):
        open_request_type_dropdown()

    for item in (arrow, text, hitbox):
        canvas.tag_bind(item, "<Enter>", enter)
        canvas.tag_bind(item, "<Leave>", leave)
        canvas.tag_bind(item, "<Button-1>", click)

    if request_type_dropdown_open:
        draw_request_type_dropdown(x, y + h // 2 + 8, w)


def draw_request_type_dropdown(x, y, w):
    global request_type_selected_index

    max_visible = 6
    row_h = 44
    visible = REQUEST_TYPE_OPTIONS[request_type_scroll:request_type_scroll + max_visible]

    box_x1 = x - w // 2
    box_x2 = x + w // 2
    box_y1 = y
    box_y2 = y + len(visible) * row_h + 14

    rounded_rect(
        box_x1,
        box_y1,
        box_x2,
        box_y2,
        r=22,
        fill="#f3eeee",
        outline="#d9d1d1",
        width=2
    )

    for index, value in enumerate(visible):
        real_index = request_type_scroll + index
        item_y = box_y1 + 18 + index * row_h + row_h // 2

        is_selected = real_index == request_type_selected_index
        if is_selected:
            selected_bg = rounded_rect(
                box_x1 + 12,
                item_y - 18,
                box_x2 - 12,
                item_y + 18,
                r=16,
                fill="#d7c28a",
                outline="#d7c28a",
                width=1
            )

        item_text = canvas.create_text(
            box_x2 - 35,
            item_y,
            text=value,
            fill="#173b38",
            font=("Arial", 17, "bold"),
            anchor="e"
        )

        hitbox = canvas.create_rectangle(
            box_x1 + 10,
            item_y - 20,
            box_x2 - 10,
            item_y + 20,
            fill="",
            outline=""
        )

        def enter(event, idx=real_index):
            global request_type_selected_index
            request_type_selected_index = idx
            root.config(cursor="hand2")

        def leave(event):
            root.config(cursor="")

        def click(event, chosen=value):
            choose_request_type(chosen)

        for item in (item_text, hitbox):
            canvas.tag_bind(item, "<Enter>", enter)
            canvas.tag_bind(item, "<Leave>", leave)
            canvas.tag_bind(item, "<Button-1>", click)

    if request_type_scroll > 0:
        canvas.create_text(x, box_y1 + 10, text="▲", fill="#777777", font=("Arial", 12, "bold"))

    if request_type_scroll + max_visible < len(REQUEST_TYPE_OPTIONS):
        canvas.create_text(x, box_y2 - 9, text="▼", fill="#777777", font=("Arial", 12, "bold"))



def set_run_font(run, size=16, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def add_doc_paragraph(doc, text, size=16, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def open_generated_file(path):
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception as e:
        try:
            from tkinter import messagebox
            messagebox.showinfo("تم", f"تم إنشاء الملف:\n{path}")
        except:
            pass


def create_job_request_word():
    output_dir = os.path.join(os.path.expanduser("~"), "IDARA_DZ_Outputs")
    os.makedirs(output_dir, exist_ok=True)

    first_name = job_form_entries.get("first_name", "").strip()
    last_name = job_form_entries.get("last_name", "").strip()
    full_name = f"{last_name} {first_name}".strip()

    request_date = job_form_entries.get("date", "").strip()
    birth_info = job_form_entries.get("birth_info", "").strip()
    address = job_form_entries.get("address", "").strip()
    phone = job_form_entries.get("phone", "").strip()
    recipient = job_form_entries.get("recipient", "").strip()
    position = job_form_entries.get("position", "").strip()
    degree = job_form_entries.get("degree", "").strip()
    specialty = job_form_entries.get("specialty", "").strip()
    request_type = job_form_entries.get("request_type", "").strip() or "طلب توظيف (عام)"

    safe_name = full_name if full_name else "بدون_اسم"
    safe_name = safe_name.replace("/", "-").replace("\\", "-").replace(":", "-")
    file_name = f"{request_type}_{safe_name}.docx".replace(" ", "_")
    output_path = os.path.join(output_dir, file_name)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.left_margin = Cm(2.0)

    # Top date
    add_doc_paragraph(
        doc,
        f"في: {request_date}",
        size=16,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=36
    )

    # Personal info block
    add_doc_paragraph(doc, f"الاسم واللقب: {full_name}", size=16, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)
    add_doc_paragraph(doc, f"تاريخ ومكان الميلاد: {birth_info}", size=16, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)
    add_doc_paragraph(doc, f"العنوان: {address}", size=16, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)
    add_doc_paragraph(doc, f"الهاتف: {phone}", size=16, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=38)

    # Recipient
    add_doc_paragraph(
        doc,
        f"إلى السيد: {recipient}",
        size=20,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=45
    )

    # Subject
    subject = "طلب توظيف"
    if request_type and request_type != "طلب توظيف (عام)":
        subject = request_type

    add_doc_paragraph(
        doc,
        f"الموضوع: {subject}",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=42
    )

    body_1 = (
        f"يشرفني أن أتقدم بطلبي هذا من أجل الحصول على منصب {position} "
        f"يتوافق مع مؤهلاتي العلمية وخبراتي المهنية في مؤسستكم الموقرة."
    )

    body_2 = (
        f"أنا {full_name}، متحصل على شهادة {degree}، تخصص {specialty}. "
        "أتمتع بروح المسؤولية والانضباط، وأطمح للمساهمة في تطوير أداء مؤسستكم."
    )

    add_doc_paragraph(doc, body_1, size=17, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=28)
    add_doc_paragraph(doc, body_2, size=17, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=80)

    add_doc_paragraph(
        doc,
        "توقيع المعني:",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=10
    )

    doc.save(output_path)
    return output_path


def make_underline_entry(x, y, w, placeholder, field_key):
    value = job_form_entries.get(field_key, "")

    line = canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b9c3bf",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 17, "bold"),
        justify="right",
        bd=0,
        bg="#173b38",
        fg="#f4f4f4",
        insertbackground="#f4f4f4"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#f4f4f4")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#8f9996")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#f4f4f4")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#8f9996")
            job_form_entries[field_key] = ""
            root.after(200, clear_suggestions)

    def save_value(event=None):
        val = entry.get()
        if val == placeholder:
            job_form_entries[field_key] = ""
            clear_suggestions()
        else:
            job_form_entries[field_key] = val
            draw_client_suggestions(x, y, w, field_key, val)

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=36)
    return entry


def draw_request_type_underline(x, y, w):
    selected_value = job_form_entries.get("request_type", "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b9c3bf",
        width=2
    )

    arrow = canvas.create_text(
        x - w // 2 + 28,
        y,
        text="▾",
        fill="#555555",
        font=("Arial", 18, "bold")
    )

    text = canvas.create_text(
        x + w // 2,
        y,
        text=selected_value if selected_value else "نوع الطلب",
        fill="#f4f4f4" if selected_value else "#8f9996",
        font=("Arial", 17, "bold"),
        anchor="e"
    )

    hitbox = canvas.create_rectangle(
        x - w // 2,
        y - 18,
        x + w // 2,
        y + 25,
        fill="",
        outline=""
    )

    def enter(event):
        root.config(cursor="hand2")

    def leave(event):
        root.config(cursor="")

    def click(event):
        open_request_type_dropdown()

    for item in (arrow, text, hitbox):
        canvas.tag_bind(item, "<Enter>", enter)
        canvas.tag_bind(item, "<Leave>", leave)
        canvas.tag_bind(item, "<Button-1>", click)

    if request_type_dropdown_open:
        draw_request_type_dropdown(x, y + 28, w)



written_request_items = [
    "طلب توظيف 1","طلب توظيف 2","مسابقة الجمارك","مسابقة الشرطة",
    "مسابقة الحماية المدنية","عقود ماقبل التشغيل","مسابقة الماستر","مسابقة على اساس الشهادة",
    "مسابقة على اساس الشهادة","مسابقة على اساس الاختبار","طلب استخلاف","طلب توظيف عام",
    "طلب استقالة","طلب سكن","طلب تسوية وضعية","طلب تحويل اداري",
    "طلب شهادة عمل","طلب بطاقة فلاح","طلب انضمام صفوف ج.و.ش","طلب عداد كهربائي",
    "طلب عداد غاز","طلب تصحيح عقد زواج","طلب تصحيح عقد ميلاد","طلب إعانة ريفية"
]

written_request_menu_scroll = 0
written_request_search_query = ""

def show_written_request_menu():
    global current_page, written_request_menu_scroll, written_request_search_query
    current_page = "written_request_menu"
    clear_screen()

    init_dynamic_database()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#efefef", outline="#efefef")
    draw_home_sidebar("home")

    search_x = width - 290
    search_w = 260
    canvas.create_line(search_x - search_w, 104, search_x + 10, 104, fill="#bdbdbd", width=2)

    search_entry = tk.Entry(
        root,
        bd=0,
        bg="#efefef",
        fg="#111111",
        font=("Arial", 18, "bold"),
        justify="right",
        insertbackground="#111111"
    )

    if written_request_search_query:
        search_entry.insert(0, written_request_search_query)
        search_entry.config(fg="#111111")
    else:
        search_entry.insert(0, "بحث")
        search_entry.config(fg="#9b9b9b")

    def s_in(e):
        if search_entry.get() == "بحث":
            search_entry.delete(0, "end")
            search_entry.config(fg="#111111")

    def s_out(e):
        global written_request_search_query
        if not search_entry.get():
            written_request_search_query = ""
            search_entry.insert(0, "بحث")
            search_entry.config(fg="#9b9b9b")

    def do_search(e=None):
        global written_request_search_query, written_request_menu_scroll
        text = search_entry.get().strip()
        if text == "بحث":
            text = ""
        written_request_search_query = text
        written_request_menu_scroll = 0
        show_written_request_menu()

    search_entry.bind("<FocusIn>", s_in)
    search_entry.bind("<FocusOut>", s_out)
    search_entry.bind("<KeyRelease>", do_search)

    canvas.create_window(search_x - 118, 82, window=search_entry, width=search_w - 42, height=30)
    canvas.create_text(search_x + 5, 82, text="⌕", fill="#a5a5a5", font=("Arial", 28))

    add_center_x = search_x - search_w - 45
    add_circle = canvas.create_oval(
        add_center_x - 24,
        58,
        add_center_x + 24,
        106,
        fill="#ffffff",
        outline="#d0d0d0",
        width=2
    )
    add_text = canvas.create_text(add_center_x, 82, text="+", fill="#000000", font=("Arial", 26, "bold"))

    def add_enter(event):
        canvas.itemconfig(add_circle, fill="#fafafa")
        root.config(cursor="hand2")

    def add_leave(event):
        canvas.itemconfig(add_circle, fill="#ffffff")
        root.config(cursor="")

    def add_click(event):
        show_dynamic_card_editor()

    for item in (add_circle, add_text):
        canvas.tag_bind(item, "<Enter>", add_enter)
        canvas.tag_bind(item, "<Leave>", add_leave)
        canvas.tag_bind(item, "<Button-1>", add_click)

    dynamic_titles = [row["title"] for row in get_dynamic_cards()]
    all_items = list(written_request_items) + dynamic_titles

    query = written_request_search_query.strip().lower()
    if query:
        filtered_items = [item for item in all_items if query in item.lower()]
    else:
        filtered_items = list(all_items)

    visible_capacity = 24
    max_scroll = max(0, len(filtered_items) - visible_capacity)
    written_request_menu_scroll = max(0, min(written_request_menu_scroll, max_scroll))
    visible_items = filtered_items[written_request_menu_scroll:written_request_menu_scroll + visible_capacity]

    columns = [
        visible_items[16:24],
        visible_items[8:16],
        visible_items[0:8],
    ]

    col_x = [315, 755, 1195]
    start_y = 175
    gap_y = 77
    card_w = 385
    card_h = 58

    for c, items in enumerate(columns):
        for r, title in enumerate(items):
            x = col_x[c]
            y = start_y + r * gap_y

            rounded_home_rect(
                x - card_w//2 + 8,
                y - card_h//2 + 10,
                x + card_w//2 + 8,
                y + card_h//2 + 10,
                r=8,
                fill="#cfcfcf",
                outline="#cfcfcf"
            )

            card = rounded_home_rect(
                x - card_w//2,
                y - card_h//2,
                x + card_w//2,
                y + card_h//2,
                r=8,
                fill="#ffffff",
                outline="#ececec"
            )

            size = 20
            sub = ""

            if title == "طلب توظيف عام":
                size = 18
                sub = "خارج إطار المسابقات"

            txt = canvas.create_text(x, y - 4, text=title, fill="#000000", font=("Arial", size, "bold"))

            subtxt = None
            if sub:
                subtxt = canvas.create_text(x, y + 16, text=sub, fill="#555555", font=("Arial", 10, "bold"))

            def enter(e, cd=card):
                canvas.itemconfig(cd, fill="#fafafa")
                root.config(cursor="hand2")

            def leave(e, cd=card):
                canvas.itemconfig(cd, fill="#ffffff")
                root.config(cursor="")

            def click(e, t=title):
                clear_dynamic_context_menu()
                job_form_entries["request_type"] = t
                if t == "طلب توظيف 1":
                    show_job_request_1_form()
                elif t == "طلب توظيف 2":
                    show_job_request_2_form()
                elif t == "مسابقة الجمارك":
                    show_customs_exam_form()
                elif t == "مسابقة الشرطة":
                    show_police_exam_form()
                elif t == "مسابقة الحماية المدنية":
                    show_civil_protection_exam_form()
                elif t == "عقود ماقبل التشغيل":
                    show_pre_employment_contract_form()
                elif t == "مسابقة الماستر":
                    show_master_exam_form()
                else:
                    dynamic_card_rows = db_fetchall("SELECT id FROM dynamic_cards WHERE title=? AND is_visible=1", (t,))
                    if dynamic_card_rows:
                        show_dynamic_card_form(dynamic_card_rows[0]["id"])
                    else:
                        show_job_request_form()

            def right_click(e, t=title):
                show_card_context_menu(e.x, e.y, t)

            bind_items = [card, txt]
            if subtxt:
                bind_items.append(subtxt)

            for it in bind_items:
                canvas.tag_bind(it, "<Enter>", enter)
                canvas.tag_bind(it, "<Leave>", leave)
                canvas.tag_bind(it, "<Button-1>", click)
                canvas.tag_bind(it, "<Button-3>", right_click)

    if not visible_items:
        canvas.create_text(760, 395, text="لا توجد نتائج", fill="#777777", font=("Arial", 28, "bold"))

    if written_request_menu_scroll > 0:
        canvas.create_text(width - 55, 145, text="▲", fill="#777777", font=("Arial", 20, "bold"))

    if written_request_menu_scroll < max_scroll:
        canvas.create_text(width - 55, height - 45, text="▼", fill="#777777", font=("Arial", 20, "bold"))


def scroll_written_request_menu(event):
    global written_request_menu_scroll

    if current_page != "written_request_menu":
        return

    dynamic_titles = [row["title"] for row in get_dynamic_cards()]
    all_items = list(written_request_items) + dynamic_titles

    query = written_request_search_query.strip().lower()
    if query:
        filtered_items = [item for item in all_items if query in item.lower()]
    else:
        filtered_items = list(all_items)

    max_scroll = max(0, len(filtered_items) - 24)

    if event.delta < 0:
        written_request_menu_scroll += 3
    else:
        written_request_menu_scroll -= 3

    written_request_menu_scroll = max(0, min(written_request_menu_scroll, max_scroll))
    show_written_request_menu()


def make_job1_entry(x, y, w, placeholder, field_key):
    value = job_request_1_entries.get(field_key, "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b8b8b8",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            job_request_1_entries[field_key] = ""

    def save_value(event=None):
        val = entry.get()
        if val == placeholder:
            job_request_1_entries[field_key] = ""
        else:
            job_request_1_entries[field_key] = val

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_job_request_1_form():
    global current_page
    current_page = "job_request_1_form"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)

    start_y = int(height * 0.12)
    gap = int(height * 0.112)

    right_fields = [
        ("الاسم", "first_name"),
        ("اللقب", "last_name"),
        ("رقم الهاتف", "phone"),
        ("العنوان", "address"),
    ]

    left_fields = [
        ("تاريخ الطلب", "request_date"),
        ("السيد / الجهة المستقبلة", "recipient"),
        ("المنصب", "position"),
        ("مدة الخبرة", "experience_duration"),
        ("الشهادة", "degree"),
    ]

    for index, (placeholder, key) in enumerate(right_fields):
        make_job1_entry(right_col_x, start_y + index * gap, field_w, placeholder, key)

    for index, (placeholder, key) in enumerate(left_fields):
        y = start_y + index * gap

        if key == "request_date":
            make_job1_entry(left_col_x + 24, y, field_w - 58, placeholder, key)

            cal_icon = canvas.create_text(
                left_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            cal_hitbox = canvas.create_rectangle(
                left_col_x - field_w // 2,
                y - 22,
                left_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_job1_cal(event):
                show_job1_calendar_picker()

            for item in (cal_icon, cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_job1_cal)

        else:
            make_job1_entry(left_col_x, y, field_w, placeholder, key)

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        create_job_request_1_word()

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)



job1_calendar_month = date.today().month
job1_calendar_year = date.today().year


def show_job1_calendar_picker():
    global current_page
    current_page = "job1_calendar_picker"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    panel_x1 = int(width * 0.24)
    panel_x2 = int(width * 0.82)
    panel_y1 = int(height * 0.12)
    panel_y2 = int(height * 0.82)

    rounded_home_rect(
        panel_x1,
        panel_y1,
        panel_x2,
        panel_y2,
        r=18,
        fill="#ffffff",
        outline="#dddddd",
        width=2
    )

    month_name = calendar.month_name[job1_calendar_month]

    canvas.create_text(
        width // 2,
        panel_y1 + 55,
        text=f"{month_name} {job1_calendar_year}",
        fill="#000000",
        font=("Arial", 28, "bold")
    )

    prev_btn = canvas.create_text(panel_x1 + 70, panel_y1 + 55, text="‹", fill="#000000", font=("Arial", 42, "bold"))
    next_btn = canvas.create_text(panel_x2 - 70, panel_y1 + 55, text="›", fill="#000000", font=("Arial", 42, "bold"))

    def prev_month(event):
        global job1_calendar_month, job1_calendar_year
        job1_calendar_month -= 1
        if job1_calendar_month < 1:
            job1_calendar_month = 12
            job1_calendar_year -= 1
        show_job1_calendar_picker()

    def next_month(event):
        global job1_calendar_month, job1_calendar_year
        job1_calendar_month += 1
        if job1_calendar_month > 12:
            job1_calendar_month = 1
            job1_calendar_year += 1
        show_job1_calendar_picker()

    for item, cmd in [(prev_btn, prev_month), (next_btn, next_month)]:
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", cmd)

    days_header = ["Sat", "Sun", "Mon", "Tue", "Wed", "Thu", "Fri"]
    grid_x1 = panel_x1 + 80
    grid_x2 = panel_x2 - 80
    col_w = (grid_x2 - grid_x1) // 7
    start_y = panel_y1 + 120
    row_h = 58

    for i, d in enumerate(days_header):
        canvas.create_text(
            grid_x1 + i * col_w + col_w // 2,
            start_y,
            text=d,
            fill="#555555",
            font=("Arial", 14, "bold")
        )

    cal = calendar.Calendar(firstweekday=5)
    days = list(cal.itermonthdays(job1_calendar_year, job1_calendar_month))

    for index, day in enumerate(days):
        row = index // 7
        col = index % 7
        x = grid_x1 + col * col_w + col_w // 2
        y = start_y + 45 + row * row_h

        if day == 0:
            continue

        day_box = rounded_home_rect(
            x - 22,
            y - 20,
            x + 22,
            y + 20,
            r=10,
            fill="#ffffff",
            outline="#dddddd",
            width=1
        )

        day_text = canvas.create_text(
            x,
            y,
            text=str(day),
            fill="#000000",
            font=("Arial", 15, "bold")
        )

        def choose_day(event, selected_day=day):
            job_request_1_entries["request_date"] = f"{selected_day:02d}/{job1_calendar_month:02d}/{job1_calendar_year}"
            show_job_request_1_form()

        for item in (day_box, day_text):
            canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
            canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
            canvas.tag_bind(item, "<Button-1>", choose_day)

def set_job1_run_font(run, size=14, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def add_job1_paragraph(doc, text="", size=14, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_job1_run_font(run, size=size, bold=bold)
    return p


def create_job_request_1_word():
    output_dir = os.path.join(os.path.expanduser("~"), "IDARA_DZ_Outputs")
    os.makedirs(output_dir, exist_ok=True)

    first_name = job_request_1_entries.get("first_name", "").strip()
    last_name = job_request_1_entries.get("last_name", "").strip()
    phone = job_request_1_entries.get("phone", "").strip()
    address = job_request_1_entries.get("address", "").strip()
    recipient = job_request_1_entries.get("recipient", "").strip()
    position = job_request_1_entries.get("position", "").strip()
    experience_duration = job_request_1_entries.get("experience_duration", "").strip()
    degree = job_request_1_entries.get("degree", "").strip()
    request_date = job_request_1_entries.get("request_date", "").strip()

    full_name = f"{last_name} {first_name}".strip()
    safe_name = full_name if full_name else "بدون_اسم"
    safe_name = safe_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    output_path = os.path.join(output_dir, f"طلب_توظيف_1_{safe_name}.docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.right_margin = Cm(1.5)
    section.left_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(f"التاريخ: {request_date if request_date else '-- / -- / 2026'}.")
    set_job1_run_font(r, 13, True)

    add_job1_paragraph(doc, f"الاسم: {first_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"اللقب: {last_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"الهاتف: {phone}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"العنوان: {address}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 28)

    add_job1_paragraph(doc, f"الى السيد: {recipient}", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 60)

    add_job1_paragraph(doc, "الموضوع:", 14, True, WD_ALIGN_PARAGRAPH.RIGHT, 8)
    add_job1_paragraph(doc, "طلب توظيف", 14, True, WD_ALIGN_PARAGRAPH.CENTER, 38)

    body = (
        f"لي عظيم الشرف أن أتقدم إلى سيادتكم بطلبي هذا والمتمثل في طلب توظيف في منصب {position}، "
        f"واحيطكم علما أني مستوفي لجميع الشروط المطلوبة ومن بينها شهادة و خبرة مهنية لمدة {experience_duration} "
        f"في المجال {degree}، وتجدون التوضيح المفصل في السيرة الذاتية الملحقة مع هذا الطلب."
    )

    add_job1_paragraph(doc, body, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 22)
    add_job1_paragraph(doc, "في انتظار ردكم تقبلوا منا سيدي فائق التقدير والاحترام.", 13, False, WD_ALIGN_PARAGRAPH.CENTER, 120)

    add_job1_paragraph(doc, "امضاء المعني", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 8)

    doc.save(output_path)

    try:
        if sys.platform.startswith("win"):
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", output_path])
        else:
            subprocess.Popen(["xdg-open", output_path])
    except Exception:
        pass




job_request_2_entries = {}
job2_calendar_month = date.today().month
job2_calendar_year = date.today().year


def make_job2_entry(x, y, w, placeholder, field_key):
    value = job_request_2_entries.get(field_key, "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b8b8b8",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            job_request_2_entries[field_key] = ""

    def save_value(event=None):
        val = entry.get()
        if val == placeholder:
            job_request_2_entries[field_key] = ""
        else:
            job_request_2_entries[field_key] = val

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_job_request_2_form():
    global current_page
    current_page = "job_request_2_form"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)

    start_y = int(height * 0.12)
    gap = int(height * 0.095)

    right_fields = [
        ("الاسم", "first_name"),
        ("اللقب", "last_name"),
        ("رقم الهاتف", "phone"),
        ("العنوان", "address"),
        ("تاريخ الطلب", "request_date"),
    ]

    left_fields = [
        ("الى السيد / الجهة المستقبلة", "recipient"),
        ("الشهادة", "degree"),
        ("الجامعة", "university"),
        ("ولاية الجامعة", "university_state"),
    ]

    for index, (placeholder, key) in enumerate(right_fields):
        y = start_y + index * gap

        if key == "request_date":
            make_job2_entry(right_col_x + 24, y, field_w - 58, placeholder, key)

            cal_icon = canvas.create_text(
                right_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            cal_hitbox = canvas.create_rectangle(
                right_col_x - field_w // 2,
                y - 22,
                right_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_job2_cal(event):
                show_job2_calendar_picker()

            for item in (cal_icon, cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_job2_cal)
        else:
            make_job2_entry(right_col_x, y, field_w, placeholder, key)

    for index, (placeholder, key) in enumerate(left_fields):
        make_job2_entry(left_col_x, start_y + index * gap, field_w, placeholder, key)

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        create_job_request_2_word()

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def show_job2_calendar_picker():
    global current_page
    current_page = "job2_calendar_picker"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    panel_x1 = int(width * 0.24)
    panel_x2 = int(width * 0.82)
    panel_y1 = int(height * 0.12)
    panel_y2 = int(height * 0.82)

    rounded_home_rect(
        panel_x1,
        panel_y1,
        panel_x2,
        panel_y2,
        r=18,
        fill="#ffffff",
        outline="#dddddd",
        width=2
    )

    month_name = calendar.month_name[job2_calendar_month]

    canvas.create_text(
        width // 2,
        panel_y1 + 55,
        text=f"{month_name} {job2_calendar_year}",
        fill="#000000",
        font=("Arial", 28, "bold")
    )

    prev_btn = canvas.create_text(panel_x1 + 70, panel_y1 + 55, text="‹", fill="#000000", font=("Arial", 42, "bold"))
    next_btn = canvas.create_text(panel_x2 - 70, panel_y1 + 55, text="›", fill="#000000", font=("Arial", 42, "bold"))

    def prev_month(event):
        global job2_calendar_month, job2_calendar_year
        job2_calendar_month -= 1
        if job2_calendar_month < 1:
            job2_calendar_month = 12
            job2_calendar_year -= 1
        show_job2_calendar_picker()

    def next_month(event):
        global job2_calendar_month, job2_calendar_year
        job2_calendar_month += 1
        if job2_calendar_month > 12:
            job2_calendar_month = 1
            job2_calendar_year += 1
        show_job2_calendar_picker()

    for item, cmd in [(prev_btn, prev_month), (next_btn, next_month)]:
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", cmd)

    days_header = ["Sat", "Sun", "Mon", "Tue", "Wed", "Thu", "Fri"]
    grid_x1 = panel_x1 + 80
    grid_x2 = panel_x2 - 80
    col_w = (grid_x2 - grid_x1) // 7
    start_y = panel_y1 + 120
    row_h = 58

    for i, d in enumerate(days_header):
        canvas.create_text(
            grid_x1 + i * col_w + col_w // 2,
            start_y,
            text=d,
            fill="#555555",
            font=("Arial", 14, "bold")
        )

    cal = calendar.Calendar(firstweekday=5)
    days = list(cal.itermonthdays(job2_calendar_year, job2_calendar_month))

    for index, day in enumerate(days):
        row = index // 7
        col = index % 7
        x = grid_x1 + col * col_w + col_w // 2
        y = start_y + 45 + row * row_h

        if day == 0:
            continue

        day_box = rounded_home_rect(
            x - 22,
            y - 20,
            x + 22,
            y + 20,
            r=10,
            fill="#ffffff",
            outline="#dddddd",
            width=1
        )

        day_text = canvas.create_text(
            x,
            y,
            text=str(day),
            fill="#000000",
            font=("Arial", 15, "bold")
        )

        def choose_day(event, selected_day=day):
            job_request_2_entries["request_date"] = f"{selected_day:02d}/{job2_calendar_month:02d}/{job2_calendar_year}"
            show_job_request_2_form()

        for item in (day_box, day_text):
            canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
            canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
            canvas.tag_bind(item, "<Button-1>", choose_day)


def create_job_request_2_word():
    output_dir = os.path.join(os.path.expanduser("~"), "IDARA_DZ_Outputs")
    os.makedirs(output_dir, exist_ok=True)

    first_name = job_request_2_entries.get("first_name", "").strip()
    last_name = job_request_2_entries.get("last_name", "").strip()
    phone = job_request_2_entries.get("phone", "").strip()
    address = job_request_2_entries.get("address", "").strip()
    request_date = job_request_2_entries.get("request_date", "").strip()
    recipient = job_request_2_entries.get("recipient", "").strip()
    degree = job_request_2_entries.get("degree", "").strip()
    university = job_request_2_entries.get("university", "").strip()
    university_state = job_request_2_entries.get("university_state", "").strip()

    full_name = f"{last_name} {first_name}".strip()
    safe_name = full_name if full_name else "بدون_اسم"
    safe_name = safe_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    output_path = os.path.join(output_dir, f"طلب_توظيف_2_{safe_name}.docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.right_margin = Cm(1.5)
    section.left_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(f"التاريخ: {request_date if request_date else '-- / -- / 2026'}.")
    set_job1_run_font(r, 13, True)

    add_job1_paragraph(doc, f"الاسم: {first_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"اللقب: {last_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"الهاتف: {phone}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"العنوان: {address}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 55)

    add_job1_paragraph(doc, f"الى السيد: {recipient}", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 55)

    add_job1_paragraph(doc, "الموضوع:", 14, True, WD_ALIGN_PARAGRAPH.RIGHT, 8)
    add_job1_paragraph(doc, "طلب توظيف", 14, True, WD_ALIGN_PARAGRAPH.CENTER, 38)

    body = (
        f"لي عظيم الشرف أن أتقدم إلى سيادتكم بطلبي هذا والمتمثل في طلب وظيفة "
        f"وعلما أني متحصل على شهادة {degree} من طرف جامعة {university} ولاية {university_state}. "
        "كما أني معفى من جميع التزامات الخدمة الوطنية."
    )

    add_job1_paragraph(doc, body, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 28)
    add_job1_paragraph(doc, "في انتظار ردكم تقبلوا منا سيدي فائق التقدير والاحترام.", 13, False, WD_ALIGN_PARAGRAPH.CENTER, 120)

    add_job1_paragraph(doc, "امضاء المعني", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 8)

    doc.save(output_path)

    try:
        if sys.platform.startswith("win"):
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", output_path])
        else:
            subprocess.Popen(["xdg-open", output_path])
    except Exception:
        pass


customs_exam_entries = {}
customs_calendar_month = date.today().month
customs_calendar_year = date.today().year
customs_birth_calendar_month = date.today().month
customs_birth_calendar_year = date.today().year


def make_customs_entry(x, y, w, placeholder, field_key):
    value = customs_exam_entries.get(field_key, "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b8b8b8",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            customs_exam_entries[field_key] = ""

    def save_value(event=None):
        val = entry.get()
        if val == placeholder:
            customs_exam_entries[field_key] = ""
        else:
            customs_exam_entries[field_key] = val

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_customs_exam_form():
    global current_page
    current_page = "customs_exam_form"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)

    start_y = int(height * 0.10)
    gap = int(height * 0.085)

    right_fields = [
        ("الاسم", "first_name"),
        ("اللقب", "last_name"),
        ("تاريخ الميلاد", "birth_date"),
        ("العنوان", "address"),
        ("رقم الهاتف", "phone"),
        ("تاريخ الطلب", "request_date"),
    ]

    left_fields = [
        ("المديرية الجهوية للجمارك لولاية", "customs_regional_directorate"),
        ("الشهادة", "degree"),
        ("الجامعة", "university"),
        ("ولاية الجامعة", "university_state"),
    ]

    for index, (placeholder, key) in enumerate(right_fields):
        y = start_y + index * gap

        if key == "request_date":
            make_customs_entry(right_col_x + 24, y, field_w - 58, placeholder, key)

            cal_icon = canvas.create_text(
                right_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            cal_hitbox = canvas.create_rectangle(
                right_col_x - field_w // 2,
                y - 22,
                right_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_customs_cal(event):
                show_customs_calendar_picker()

            for item in (cal_icon, cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_customs_cal)

        elif key == "birth_date":
            make_customs_entry(right_col_x + 24, y, field_w - 58, placeholder, key)

            birth_cal_icon = canvas.create_text(
                right_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            birth_cal_hitbox = canvas.create_rectangle(
                right_col_x - field_w // 2,
                y - 22,
                right_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_birth_cal(event):
                show_customs_birth_calendar_picker()

            for item in (birth_cal_icon, birth_cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_birth_cal)

        else:
            make_customs_entry(right_col_x, y, field_w, placeholder, key)

    for index, (placeholder, key) in enumerate(left_fields):
        make_customs_entry(left_col_x, start_y + index * gap, field_w, placeholder, key)

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        create_customs_exam_word()

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def show_customs_calendar_picker():
    global current_page
    current_page = "customs_calendar_picker"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    panel_x1 = int(width * 0.24)
    panel_x2 = int(width * 0.82)
    panel_y1 = int(height * 0.12)
    panel_y2 = int(height * 0.82)

    rounded_home_rect(panel_x1, panel_y1, panel_x2, panel_y2, r=18, fill="#ffffff", outline="#dddddd", width=2)

    month_name = calendar.month_name[customs_calendar_month]

    canvas.create_text(
        width // 2,
        panel_y1 + 55,
        text=f"{month_name} {customs_calendar_year}",
        fill="#000000",
        font=("Arial", 28, "bold")
    )

    prev_btn = canvas.create_text(panel_x1 + 70, panel_y1 + 55, text="‹", fill="#000000", font=("Arial", 42, "bold"))
    next_btn = canvas.create_text(panel_x2 - 70, panel_y1 + 55, text="›", fill="#000000", font=("Arial", 42, "bold"))

    def prev_month(event):
        global customs_calendar_month, customs_calendar_year
        customs_calendar_month -= 1
        if customs_calendar_month < 1:
            customs_calendar_month = 12
            customs_calendar_year -= 1
        show_customs_calendar_picker()

    def next_month(event):
        global customs_calendar_month, customs_calendar_year
        customs_calendar_month += 1
        if customs_calendar_month > 12:
            customs_calendar_month = 1
            customs_calendar_year += 1
        show_customs_calendar_picker()

    for item, cmd in [(prev_btn, prev_month), (next_btn, next_month)]:
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", cmd)

    days_header = ["Sat", "Sun", "Mon", "Tue", "Wed", "Thu", "Fri"]
    grid_x1 = panel_x1 + 80
    grid_x2 = panel_x2 - 80
    col_w = (grid_x2 - grid_x1) // 7
    start_y = panel_y1 + 120
    row_h = 58

    for i, d in enumerate(days_header):
        canvas.create_text(grid_x1 + i * col_w + col_w // 2, start_y, text=d, fill="#555555", font=("Arial", 14, "bold"))

    cal = calendar.Calendar(firstweekday=5)
    days = list(cal.itermonthdays(customs_calendar_year, customs_calendar_month))

    for index, day in enumerate(days):
        row = index // 7
        col = index % 7
        x = grid_x1 + col * col_w + col_w // 2
        y = start_y + 45 + row * row_h

        if day == 0:
            continue

        day_box = rounded_home_rect(x - 22, y - 20, x + 22, y + 20, r=10, fill="#ffffff", outline="#dddddd", width=1)
        day_text = canvas.create_text(x, y, text=str(day), fill="#000000", font=("Arial", 15, "bold"))

        def choose_day(event, selected_day=day):
            customs_exam_entries["request_date"] = f"{selected_day:02d}/{customs_calendar_month:02d}/{customs_calendar_year}"
            show_customs_exam_form()

        for item in (day_box, day_text):
            canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
            canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
            canvas.tag_bind(item, "<Button-1>", choose_day)


def show_customs_birth_calendar_picker():
    global current_page
    current_page = "customs_birth_calendar_picker"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    panel_x1 = int(width * 0.24)
    panel_x2 = int(width * 0.82)
    panel_y1 = int(height * 0.12)
    panel_y2 = int(height * 0.82)

    rounded_home_rect(panel_x1, panel_y1, panel_x2, panel_y2, r=18, fill="#ffffff", outline="#dddddd", width=2)

    month_name = calendar.month_name[customs_birth_calendar_month]

    canvas.create_text(
        width // 2,
        panel_y1 + 55,
        text=f"{month_name} {customs_birth_calendar_year}",
        fill="#000000",
        font=("Arial", 28, "bold")
    )

    prev_btn = canvas.create_text(panel_x1 + 70, panel_y1 + 55, text="‹", fill="#000000", font=("Arial", 42, "bold"))
    next_btn = canvas.create_text(panel_x2 - 70, panel_y1 + 55, text="›", fill="#000000", font=("Arial", 42, "bold"))

    def prev_month(event):
        global customs_birth_calendar_month, customs_birth_calendar_year
        customs_birth_calendar_month -= 1
        if customs_birth_calendar_month < 1:
            customs_birth_calendar_month = 12
            customs_birth_calendar_year -= 1
        show_customs_birth_calendar_picker()

    def next_month(event):
        global customs_birth_calendar_month, customs_birth_calendar_year
        customs_birth_calendar_month += 1
        if customs_birth_calendar_month > 12:
            customs_birth_calendar_month = 1
            customs_birth_calendar_year += 1
        show_customs_birth_calendar_picker()

    for item, cmd in [(prev_btn, prev_month), (next_btn, next_month)]:
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", cmd)

    days_header = ["Sat", "Sun", "Mon", "Tue", "Wed", "Thu", "Fri"]
    grid_x1 = panel_x1 + 80
    grid_x2 = panel_x2 - 80
    col_w = (grid_x2 - grid_x1) // 7
    start_y = panel_y1 + 120
    row_h = 58

    for i, d in enumerate(days_header):
        canvas.create_text(grid_x1 + i * col_w + col_w // 2, start_y, text=d, fill="#555555", font=("Arial", 14, "bold"))

    cal = calendar.Calendar(firstweekday=5)
    days = list(cal.itermonthdays(customs_birth_calendar_year, customs_birth_calendar_month))

    for index, day in enumerate(days):
        row = index // 7
        col = index % 7
        x = grid_x1 + col * col_w + col_w // 2
        y = start_y + 45 + row * row_h

        if day == 0:
            continue

        day_box = rounded_home_rect(x - 22, y - 20, x + 22, y + 20, r=10, fill="#ffffff", outline="#dddddd", width=1)
        day_text = canvas.create_text(x, y, text=str(day), fill="#000000", font=("Arial", 15, "bold"))

        def choose_day(event, selected_day=day):
            customs_exam_entries["birth_date"] = f"{selected_day:02d}/{customs_birth_calendar_month:02d}/{customs_birth_calendar_year}"
            show_customs_exam_form()

        for item in (day_box, day_text):
            canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
            canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
            canvas.tag_bind(item, "<Button-1>", choose_day)


def create_customs_exam_word():
    output_dir = os.path.join(os.path.expanduser("~"), "IDARA_DZ_Outputs")
    os.makedirs(output_dir, exist_ok=True)

    first_name = customs_exam_entries.get("first_name", "").strip()
    last_name = customs_exam_entries.get("last_name", "").strip()
    birth_date = customs_exam_entries.get("birth_date", "").strip()
    address = customs_exam_entries.get("address", "").strip()
    phone = customs_exam_entries.get("phone", "").strip()
    request_date = customs_exam_entries.get("request_date", "").strip()
    customs_state = customs_exam_entries.get("customs_regional_directorate", "").strip()
    degree = customs_exam_entries.get("degree", "").strip()
    university = customs_exam_entries.get("university", "").strip()
    university_state = customs_exam_entries.get("university_state", "").strip()

    full_name = f"{last_name} {first_name}".strip()
    safe_name = full_name if full_name else "بدون_اسم"
    safe_name = safe_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    output_path = os.path.join(output_dir, f"مسابقة_الجمارك_{safe_name}.docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.right_margin = Cm(1.5)
    section.left_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(f"التاريخ: {request_date if request_date else '-- / -- / 2026'}.")
    set_job1_run_font(r, 13, True)

    add_job1_paragraph(doc, f"الاسم: {first_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"اللقب: {last_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"الهاتف: {phone}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"العنوان: {address}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"تاريخ الميلاد: {birth_date}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 18)

    add_job1_paragraph(doc, "المديرية العامة للجمارك", 13, False, WD_ALIGN_PARAGRAPH.LEFT, 12)
    add_job1_paragraph(doc, f"المديرية الجهوية للجمارك لولاية {customs_state}", 13, False, WD_ALIGN_PARAGRAPH.LEFT, 55)

    add_job1_paragraph(doc, "الموضوع:", 14, True, WD_ALIGN_PARAGRAPH.RIGHT, 8)
    add_job1_paragraph(
        doc,
        "طلب المشاركة في مسابقة التوظيف الخارجية على أساس الاختبار لرتبة - ضباط الفرق -",
        13,
        True,
        WD_ALIGN_PARAGRAPH.CENTER,
        36
    )

    body_1 = (
        "لي عظيم الشرف أن أتقدم إلى سيادتكم بطلبي هذا والمتمثل في طلب المشاركة في مسابقة التوظيف الخارجية على أساس "
        "الاختبار لرتبة أعوان الرقابة للجمارك ذكور / إناث لسنة ......."
    )

    body_2 = (
        f"علما أني متحصل على شهادة {degree} من طرف جامعة {university} لولاية {university_state}."
    )

    body_3 = "و معفى من جميع التزامات الخدمة الوطنية و مستعد للعمل في كافة أنحاء الوطن و في كل الظروف."

    add_job1_paragraph(doc, body_1, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 16)
    add_job1_paragraph(doc, body_2, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 16)
    add_job1_paragraph(doc, body_3, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 20)
    add_job1_paragraph(doc, "في انتظار ردكم تقبلوا منا سيدي فائق التقدير والاحترام.", 13, False, WD_ALIGN_PARAGRAPH.CENTER, 115)

    add_job1_paragraph(doc, "امضاء المعني", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 8)

    doc.save(output_path)

    try:
        if sys.platform.startswith("win"):
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", output_path])
        else:
            subprocess.Popen(["xdg-open", output_path])
    except Exception:
        pass


police_exam_entries = {}
police_calendar_month = date.today().month
police_calendar_year = date.today().year
police_birth_calendar_month = date.today().month
police_birth_calendar_year = date.today().year


def make_police_entry(x, y, w, placeholder, field_key):
    value = police_exam_entries.get(field_key, "")

    canvas.create_line(x - w // 2, y + 18, x + w // 2, y + 18, fill="#b8b8b8", width=2)

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            police_exam_entries[field_key] = ""

    def save_value(event=None):
        val = entry.get()
        police_exam_entries[field_key] = "" if val == placeholder else val

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_police_exam_form():
    global current_page
    current_page = "police_exam_form"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)

    start_y = int(height * 0.08)
    gap = int(height * 0.075)

    right_fields = [
        ("الاسم", "first_name"),
        ("اللقب", "last_name"),
        ("تاريخ الميلاد", "birth_date"),
        ("العنوان", "address"),
        ("رقم الهاتف", "phone"),
        ("تاريخ الطلب", "request_date"),
    ]

    left_fields = [
        ("الى السيد / الجهة المستقبلة", "recipient"),
        ("الرتبة", "rank"),
        ("الشهادة", "degree"),
        ("التخصص", "specialty"),
        ("الجامعة", "university"),
    ]

    for index, (placeholder, key) in enumerate(right_fields):
        y = start_y + index * gap

        if key == "request_date":
            make_police_entry(right_col_x + 24, y, field_w - 58, placeholder, key)
            cal_icon = canvas.create_text(right_col_x - field_w // 2 + 28, y, text="📅", fill="#000000", font=("Arial", 22, "bold"))
            cal_hitbox = canvas.create_rectangle(right_col_x - field_w // 2, y - 22, right_col_x - field_w // 2 + 58, y + 25, fill="", outline="")
            def open_police_cal(event):
                show_police_calendar_picker()
            for item in (cal_icon, cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_police_cal)

        elif key == "birth_date":
            make_police_entry(right_col_x + 24, y, field_w - 58, placeholder, key)
            birth_cal_icon = canvas.create_text(right_col_x - field_w // 2 + 28, y, text="📅", fill="#000000", font=("Arial", 22, "bold"))
            birth_cal_hitbox = canvas.create_rectangle(right_col_x - field_w // 2, y - 22, right_col_x - field_w // 2 + 58, y + 25, fill="", outline="")
            def open_birth_cal(event):
                show_police_birth_calendar_picker()
            for item in (birth_cal_icon, birth_cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_birth_cal)
        else:
            make_police_entry(right_col_x, y, field_w, placeholder, key)

    for index, (placeholder, key) in enumerate(left_fields):
        make_police_entry(left_col_x, start_y + index * gap, field_w, placeholder, key)

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        create_police_exam_word()

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def draw_police_calendar(kind):
    global current_page
    current_page = "police_birth_calendar_picker" if kind == "birth" else "police_calendar_picker"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    if kind == "birth":
        month = police_birth_calendar_month
        year = police_birth_calendar_year
    else:
        month = police_calendar_month
        year = police_calendar_year

    panel_x1 = int(width * 0.24)
    panel_x2 = int(width * 0.82)
    panel_y1 = int(height * 0.12)
    panel_y2 = int(height * 0.82)

    rounded_home_rect(panel_x1, panel_y1, panel_x2, panel_y2, r=18, fill="#ffffff", outline="#dddddd", width=2)

    canvas.create_text(width // 2, panel_y1 + 55, text=f"{calendar.month_name[month]} {year}", fill="#000000", font=("Arial", 28, "bold"))

    prev_btn = canvas.create_text(panel_x1 + 70, panel_y1 + 55, text="‹", fill="#000000", font=("Arial", 42, "bold"))
    next_btn = canvas.create_text(panel_x2 - 70, panel_y1 + 55, text="›", fill="#000000", font=("Arial", 42, "bold"))

    def prev_month(event):
        global police_calendar_month, police_calendar_year, police_birth_calendar_month, police_birth_calendar_year
        if kind == "birth":
            police_birth_calendar_month -= 1
            if police_birth_calendar_month < 1:
                police_birth_calendar_month = 12
                police_birth_calendar_year -= 1
            show_police_birth_calendar_picker()
        else:
            police_calendar_month -= 1
            if police_calendar_month < 1:
                police_calendar_month = 12
                police_calendar_year -= 1
            show_police_calendar_picker()

    def next_month(event):
        global police_calendar_month, police_calendar_year, police_birth_calendar_month, police_birth_calendar_year
        if kind == "birth":
            police_birth_calendar_month += 1
            if police_birth_calendar_month > 12:
                police_birth_calendar_month = 1
                police_birth_calendar_year += 1
            show_police_birth_calendar_picker()
        else:
            police_calendar_month += 1
            if police_calendar_month > 12:
                police_calendar_month = 1
                police_calendar_year += 1
            show_police_calendar_picker()

    for item, cmd in [(prev_btn, prev_month), (next_btn, next_month)]:
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", cmd)

    days_header = ["Sat", "Sun", "Mon", "Tue", "Wed", "Thu", "Fri"]
    grid_x1 = panel_x1 + 80
    grid_x2 = panel_x2 - 80
    col_w = (grid_x2 - grid_x1) // 7
    start_y = panel_y1 + 120
    row_h = 58

    for i, d in enumerate(days_header):
        canvas.create_text(grid_x1 + i * col_w + col_w // 2, start_y, text=d, fill="#555555", font=("Arial", 14, "bold"))

    cal = calendar.Calendar(firstweekday=5)
    days = list(cal.itermonthdays(year, month))

    for index, day in enumerate(days):
        row = index // 7
        col = index % 7
        x = grid_x1 + col * col_w + col_w // 2
        y = start_y + 45 + row * row_h

        if day == 0:
            continue

        day_box = rounded_home_rect(x - 22, y - 20, x + 22, y + 20, r=10, fill="#ffffff", outline="#dddddd", width=1)
        day_text = canvas.create_text(x, y, text=str(day), fill="#000000", font=("Arial", 15, "bold"))

        def choose_day(event, selected_day=day):
            if kind == "birth":
                police_exam_entries["birth_date"] = f"{selected_day:02d}/{police_birth_calendar_month:02d}/{police_birth_calendar_year}"
            else:
                police_exam_entries["request_date"] = f"{selected_day:02d}/{police_calendar_month:02d}/{police_calendar_year}"
            show_police_exam_form()

        for item in (day_box, day_text):
            canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
            canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
            canvas.tag_bind(item, "<Button-1>", choose_day)


def show_police_calendar_picker():
    draw_police_calendar("request")


def show_police_birth_calendar_picker():
    draw_police_calendar("birth")


def create_police_exam_word():
    output_dir = os.path.join(os.path.expanduser("~"), "IDARA_DZ_Outputs")
    os.makedirs(output_dir, exist_ok=True)

    first_name = police_exam_entries.get("first_name", "").strip()
    last_name = police_exam_entries.get("last_name", "").strip()
    birth_date = police_exam_entries.get("birth_date", "").strip()
    address = police_exam_entries.get("address", "").strip()
    phone = police_exam_entries.get("phone", "").strip()
    request_date = police_exam_entries.get("request_date", "").strip()
    recipient = police_exam_entries.get("recipient", "").strip()
    rank = police_exam_entries.get("rank", "").strip()
    degree = police_exam_entries.get("degree", "").strip()
    specialty = police_exam_entries.get("specialty", "").strip()
    university = police_exam_entries.get("university", "").strip()

    full_name = f"{last_name} {first_name}".strip()
    safe_name = full_name if full_name else "بدون_اسم"
    safe_name = safe_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    output_path = os.path.join(output_dir, f"مسابقة_الشرطة_{safe_name}.docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.right_margin = Cm(1.5)
    section.left_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(f"التاريخ: {request_date if request_date else '-- / -- / 2026'}.")
    set_job1_run_font(r, 13, True)

    add_job1_paragraph(doc, f"الاسم: {first_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"اللقب: {last_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"الهاتف: {phone}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"العنوان: {address}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"تاريخ الميلاد: {birth_date}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 28)

    add_job1_paragraph(doc, f"الى السيد: {recipient}", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 70)

    add_job1_paragraph(doc, "الموضوع:", 14, True, WD_ALIGN_PARAGRAPH.RIGHT, 8)
    add_job1_paragraph(doc, "طلب المشاركة في المسابقة", 14, True, WD_ALIGN_PARAGRAPH.CENTER, 38)

    body = (
        f"لي عظيم الشرف أن أتقدم إلى سيادتكم بطلبي هذا والمتمثل في طلب المشاركة في المسابقة المنظمة من طرفكم للالتحاق "
        f"برتبة {rank}، مع العلم أني متحصل على شهادة {degree} تخصص {specialty} من جامعة {university}، "
        "و معفى من كافة التزامات الخدمة الوطنية."
    )

    add_job1_paragraph(doc, body, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 22)
    add_job1_paragraph(doc, "في انتظار ردكم تقبلوا منا سيدي فائق التقدير والاحترام.", 13, False, WD_ALIGN_PARAGRAPH.CENTER, 120)
    add_job1_paragraph(doc, "امضاء المعني", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 8)

    doc.save(output_path)

    try:
        if sys.platform.startswith("win"):
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", output_path])
        else:
            subprocess.Popen(["xdg-open", output_path])
    except Exception:
        pass


civil_protection_entries = {}
civil_calendar_month = date.today().month
civil_calendar_year = date.today().year
civil_birth_calendar_month = date.today().month
civil_birth_calendar_year = date.today().year


def make_civil_entry(x, y, w, placeholder, field_key):
    value = civil_protection_entries.get(field_key, "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b8b8b8",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            civil_protection_entries[field_key] = ""

    def save_value(event=None):
        val = entry.get()
        civil_protection_entries[field_key] = "" if val == placeholder else val

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_civil_protection_exam_form():
    global current_page
    current_page = "civil_protection_exam_form"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)

    start_y = int(height * 0.08)
    gap = int(height * 0.075)

    right_fields = [
        ("الاسم", "first_name"),
        ("اللقب", "last_name"),
        ("تاريخ الميلاد", "birth_date"),
        ("العنوان", "address"),
        ("رقم الهاتف", "phone"),
        ("تاريخ الطلب", "request_date"),
    ]

    left_fields = [
        ("الى السيد / الجهة المستقبلة", "recipient"),
        ("الرتبة", "rank"),
        ("المستوى", "level"),
        ("الشهادة", "degree"),
    ]

    for index, (placeholder, key) in enumerate(right_fields):
        y = start_y + index * gap

        if key == "request_date":
            make_civil_entry(right_col_x + 24, y, field_w - 58, placeholder, key)

            cal_icon = canvas.create_text(
                right_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            cal_hitbox = canvas.create_rectangle(
                right_col_x - field_w // 2,
                y - 22,
                right_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_civil_cal(event):
                show_civil_calendar_picker()

            for item in (cal_icon, cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_civil_cal)

        elif key == "birth_date":
            make_civil_entry(right_col_x + 24, y, field_w - 58, placeholder, key)

            birth_cal_icon = canvas.create_text(
                right_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            birth_cal_hitbox = canvas.create_rectangle(
                right_col_x - field_w // 2,
                y - 22,
                right_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_birth_cal(event):
                show_civil_birth_calendar_picker()

            for item in (birth_cal_icon, birth_cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_birth_cal)

        else:
            make_civil_entry(right_col_x, y, field_w, placeholder, key)

    for index, (placeholder, key) in enumerate(left_fields):
        make_civil_entry(left_col_x, start_y + index * gap, field_w, placeholder, key)

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        create_civil_protection_exam_word()

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def draw_civil_calendar(kind):
    global current_page
    current_page = "civil_birth_calendar_picker" if kind == "birth" else "civil_calendar_picker"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    if kind == "birth":
        month = civil_birth_calendar_month
        year = civil_birth_calendar_year
    else:
        month = civil_calendar_month
        year = civil_calendar_year

    panel_x1 = int(width * 0.24)
    panel_x2 = int(width * 0.82)
    panel_y1 = int(height * 0.12)
    panel_y2 = int(height * 0.82)

    rounded_home_rect(
        panel_x1,
        panel_y1,
        panel_x2,
        panel_y2,
        r=18,
        fill="#ffffff",
        outline="#dddddd",
        width=2
    )

    canvas.create_text(
        width // 2,
        panel_y1 + 55,
        text=f"{calendar.month_name[month]} {year}",
        fill="#000000",
        font=("Arial", 28, "bold")
    )

    prev_btn = canvas.create_text(panel_x1 + 70, panel_y1 + 55, text="‹", fill="#000000", font=("Arial", 42, "bold"))
    next_btn = canvas.create_text(panel_x2 - 70, panel_y1 + 55, text="›", fill="#000000", font=("Arial", 42, "bold"))

    def prev_month(event):
        global civil_calendar_month, civil_calendar_year, civil_birth_calendar_month, civil_birth_calendar_year
        if kind == "birth":
            civil_birth_calendar_month -= 1
            if civil_birth_calendar_month < 1:
                civil_birth_calendar_month = 12
                civil_birth_calendar_year -= 1
            show_civil_birth_calendar_picker()
        else:
            civil_calendar_month -= 1
            if civil_calendar_month < 1:
                civil_calendar_month = 12
                civil_calendar_year -= 1
            show_civil_calendar_picker()

    def next_month(event):
        global civil_calendar_month, civil_calendar_year, civil_birth_calendar_month, civil_birth_calendar_year
        if kind == "birth":
            civil_birth_calendar_month += 1
            if civil_birth_calendar_month > 12:
                civil_birth_calendar_month = 1
                civil_birth_calendar_year += 1
            show_civil_birth_calendar_picker()
        else:
            civil_calendar_month += 1
            if civil_calendar_month > 12:
                civil_calendar_month = 1
                civil_calendar_year += 1
            show_civil_calendar_picker()

    for item, cmd in [(prev_btn, prev_month), (next_btn, next_month)]:
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", cmd)

    days_header = ["Sat", "Sun", "Mon", "Tue", "Wed", "Thu", "Fri"]
    grid_x1 = panel_x1 + 80
    grid_x2 = panel_x2 - 80
    col_w = (grid_x2 - grid_x1) // 7
    start_y = panel_y1 + 120
    row_h = 58

    for i, d in enumerate(days_header):
        canvas.create_text(
            grid_x1 + i * col_w + col_w // 2,
            start_y,
            text=d,
            fill="#555555",
            font=("Arial", 14, "bold")
        )

    cal = calendar.Calendar(firstweekday=5)
    days = list(cal.itermonthdays(year, month))

    for index, day in enumerate(days):
        row = index // 7
        col = index % 7
        x = grid_x1 + col * col_w + col_w // 2
        y = start_y + 45 + row * row_h

        if day == 0:
            continue

        day_box = rounded_home_rect(
            x - 22,
            y - 20,
            x + 22,
            y + 20,
            r=10,
            fill="#ffffff",
            outline="#dddddd",
            width=1
        )

        day_text = canvas.create_text(
            x,
            y,
            text=str(day),
            fill="#000000",
            font=("Arial", 15, "bold")
        )

        def choose_day(event, selected_day=day):
            if kind == "birth":
                civil_protection_entries["birth_date"] = f"{selected_day:02d}/{civil_birth_calendar_month:02d}/{civil_birth_calendar_year}"
            else:
                civil_protection_entries["request_date"] = f"{selected_day:02d}/{civil_calendar_month:02d}/{civil_calendar_year}"
            show_civil_protection_exam_form()

        for item in (day_box, day_text):
            canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
            canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
            canvas.tag_bind(item, "<Button-1>", choose_day)


def show_civil_calendar_picker():
    draw_civil_calendar("request")


def show_civil_birth_calendar_picker():
    draw_civil_calendar("birth")


def create_civil_protection_exam_word():
    output_dir = os.path.join(os.path.expanduser("~"), "IDARA_DZ_Outputs")
    os.makedirs(output_dir, exist_ok=True)

    first_name = civil_protection_entries.get("first_name", "").strip()
    last_name = civil_protection_entries.get("last_name", "").strip()
    birth_date = civil_protection_entries.get("birth_date", "").strip()
    address = civil_protection_entries.get("address", "").strip()
    phone = civil_protection_entries.get("phone", "").strip()
    request_date = civil_protection_entries.get("request_date", "").strip()
    recipient = civil_protection_entries.get("recipient", "").strip()
    rank = civil_protection_entries.get("rank", "").strip()
    level = civil_protection_entries.get("level", "").strip()
    degree = civil_protection_entries.get("degree", "").strip()

    full_name = f"{last_name} {first_name}".strip()
    safe_name = full_name if full_name else "بدون_اسم"
    safe_name = safe_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    output_path = os.path.join(output_dir, f"مسابقة_الحماية_المدنية_{safe_name}.docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.right_margin = Cm(1.5)
    section.left_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(f"التاريخ: {request_date if request_date else '-- / -- / 2026'}.")
    set_job1_run_font(r, 13, True)

    add_job1_paragraph(doc, f"الاسم: {first_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"اللقب: {last_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"الهاتف: {phone}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"العنوان: {address}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"تاريخ الميلاد: {birth_date}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 26)

    add_job1_paragraph(doc, f"الى السيد: {recipient}", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 70)

    add_job1_paragraph(doc, "الموضوع:", 14, True, WD_ALIGN_PARAGRAPH.RIGHT, 8)
    add_job1_paragraph(doc, "طلب المشاركة في المسابقة", 14, True, WD_ALIGN_PARAGRAPH.CENTER, 38)

    body_1 = (
        f"لي عظيم الشرف أن أتقدم إلى سيادتكم بطلبي هذا والمتمثل في طلب المشاركة في المسابقة المنظمة من طرفكم "
        f"للالتحاق برتبة {rank}، مع العلم أني متحصل على مستوى {level} وحامل لشهادة {degree}."
    )

    body_2 = "كما أتوفر على الشروط المطلوبة والمنصوص عليها في شروط المسابقة."

    add_job1_paragraph(doc, body_1, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 18)
    add_job1_paragraph(doc, body_2, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 22)
    add_job1_paragraph(doc, "في انتظار ردكم تقبلوا منا سيدي فائق التقدير والاحترام.", 13, False, WD_ALIGN_PARAGRAPH.CENTER, 120)

    add_job1_paragraph(doc, "امضاء المعني", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 8)

    doc.save(output_path)

    try:
        if sys.platform.startswith("win"):
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", output_path])
        else:
            subprocess.Popen(["xdg-open", output_path])
    except Exception:
        pass


pre_employment_entries = {}
pre_employment_calendar_month = date.today().month
pre_employment_calendar_year = date.today().year


def make_pre_employment_entry(x, y, w, placeholder, field_key):
    value = pre_employment_entries.get(field_key, "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b8b8b8",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            pre_employment_entries[field_key] = ""

    def save_value(event=None):
        val = entry.get()
        pre_employment_entries[field_key] = "" if val == placeholder else val

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_pre_employment_contract_form():
    global current_page
    current_page = "pre_employment_contract_form"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)

    start_y = int(height * 0.10)
    gap = int(height * 0.085)

    right_fields = [
        ("الاسم", "first_name"),
        ("اللقب", "last_name"),
        ("العنوان", "address"),
        ("رقم الهاتف", "phone"),
        ("تاريخ الطلب", "request_date"),
    ]

    left_fields = [
        ("الشهادة", "degree"),
        ("التخصص", "specialty"),
        ("الدفعة", "promotion"),
    ]

    for index, (placeholder, key) in enumerate(right_fields):
        y = start_y + index * gap

        if key == "request_date":
            make_pre_employment_entry(right_col_x + 24, y, field_w - 58, placeholder, key)

            cal_icon = canvas.create_text(
                right_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            cal_hitbox = canvas.create_rectangle(
                right_col_x - field_w // 2,
                y - 22,
                right_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_pre_employment_cal(event):
                show_pre_employment_calendar_picker()

            for item in (cal_icon, cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_pre_employment_cal)

        else:
            make_pre_employment_entry(right_col_x, y, field_w, placeholder, key)

    for index, (placeholder, key) in enumerate(left_fields):
        make_pre_employment_entry(left_col_x, start_y + index * gap, field_w, placeholder, key)

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        create_pre_employment_contract_word()

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def show_pre_employment_calendar_picker():
    global current_page
    current_page = "pre_employment_calendar_picker"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    panel_x1 = int(width * 0.24)
    panel_x2 = int(width * 0.82)
    panel_y1 = int(height * 0.12)
    panel_y2 = int(height * 0.82)

    rounded_home_rect(
        panel_x1,
        panel_y1,
        panel_x2,
        panel_y2,
        r=18,
        fill="#ffffff",
        outline="#dddddd",
        width=2
    )

    canvas.create_text(
        width // 2,
        panel_y1 + 55,
        text=f"{calendar.month_name[pre_employment_calendar_month]} {pre_employment_calendar_year}",
        fill="#000000",
        font=("Arial", 28, "bold")
    )

    prev_btn = canvas.create_text(panel_x1 + 70, panel_y1 + 55, text="‹", fill="#000000", font=("Arial", 42, "bold"))
    next_btn = canvas.create_text(panel_x2 - 70, panel_y1 + 55, text="›", fill="#000000", font=("Arial", 42, "bold"))

    def prev_month(event):
        global pre_employment_calendar_month, pre_employment_calendar_year
        pre_employment_calendar_month -= 1
        if pre_employment_calendar_month < 1:
            pre_employment_calendar_month = 12
            pre_employment_calendar_year -= 1
        show_pre_employment_calendar_picker()

    def next_month(event):
        global pre_employment_calendar_month, pre_employment_calendar_year
        pre_employment_calendar_month += 1
        if pre_employment_calendar_month > 12:
            pre_employment_calendar_month = 1
            pre_employment_calendar_year += 1
        show_pre_employment_calendar_picker()

    for item, cmd in [(prev_btn, prev_month), (next_btn, next_month)]:
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", cmd)

    days_header = ["Sat", "Sun", "Mon", "Tue", "Wed", "Thu", "Fri"]
    grid_x1 = panel_x1 + 80
    grid_x2 = panel_x2 - 80
    col_w = (grid_x2 - grid_x1) // 7
    start_y = panel_y1 + 120
    row_h = 58

    for i, d in enumerate(days_header):
        canvas.create_text(
            grid_x1 + i * col_w + col_w // 2,
            start_y,
            text=d,
            fill="#555555",
            font=("Arial", 14, "bold")
        )

    cal = calendar.Calendar(firstweekday=5)
    days = list(cal.itermonthdays(pre_employment_calendar_year, pre_employment_calendar_month))

    for index, day in enumerate(days):
        row = index // 7
        col = index % 7
        x = grid_x1 + col * col_w + col_w // 2
        y = start_y + 45 + row * row_h

        if day == 0:
            continue

        day_box = rounded_home_rect(
            x - 22,
            y - 20,
            x + 22,
            y + 20,
            r=10,
            fill="#ffffff",
            outline="#dddddd",
            width=1
        )

        day_text = canvas.create_text(
            x,
            y,
            text=str(day),
            fill="#000000",
            font=("Arial", 15, "bold")
        )

        def choose_day(event, selected_day=day):
            pre_employment_entries["request_date"] = f"{selected_day:02d}/{pre_employment_calendar_month:02d}/{pre_employment_calendar_year}"
            show_pre_employment_contract_form()

        for item in (day_box, day_text):
            canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
            canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
            canvas.tag_bind(item, "<Button-1>", choose_day)


def create_pre_employment_contract_word():
    output_dir = os.path.join(os.path.expanduser("~"), "IDARA_DZ_Outputs")
    os.makedirs(output_dir, exist_ok=True)

    first_name = pre_employment_entries.get("first_name", "").strip()
    last_name = pre_employment_entries.get("last_name", "").strip()
    address = pre_employment_entries.get("address", "").strip()
    phone = pre_employment_entries.get("phone", "").strip()
    request_date = pre_employment_entries.get("request_date", "").strip()
    degree = pre_employment_entries.get("degree", "").strip()
    specialty = pre_employment_entries.get("specialty", "").strip()
    promotion = pre_employment_entries.get("promotion", "").strip()

    full_name = f"{last_name} {first_name}".strip()
    safe_name = full_name if full_name else "بدون_اسم"
    safe_name = safe_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    output_path = os.path.join(output_dir, f"عقود_ماقبل_التشغيل_{safe_name}.docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.right_margin = Cm(1.5)
    section.left_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(f"التاريخ: {request_date if request_date else '-- / -- / 2026'}.")
    set_job1_run_font(r, 13, True)

    add_job1_paragraph(doc, f"الاسم: {first_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"اللقب: {last_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"الهاتف: {phone}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"العنوان: {address}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 42)

    add_job1_paragraph(doc, "الى السيد: مدير الوكالة الوطنية للتشغيل", 13, False, WD_ALIGN_PARAGRAPH.LEFT, 8)
    add_job1_paragraph(doc, "لولاية:", 13, False, WD_ALIGN_PARAGRAPH.LEFT, 60)

    add_job1_paragraph(doc, "الموضوع:", 14, True, WD_ALIGN_PARAGRAPH.RIGHT, 8)
    add_job1_paragraph(doc, "طلب الحصول على عمل في اطار عقود ما قبل التشغيل", 14, True, WD_ALIGN_PARAGRAPH.CENTER, 38)

    body = (
        "لي عظيم الشرف أن أتقدم إلى سيادتكم بطلبي هذا والمتمثل في طلب في الحصول على عمل في إطار عقود ما قبل التشغيل "
        f"و احيطكم علما اني متحصل على شهادة {degree} تخصص {specialty} دفعة {promotion}."
    )

    add_job1_paragraph(doc, body, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 24)
    add_job1_paragraph(doc, "في انتظار ردكم تقبلوا منا سيدي فائق التقدير والاحترام.", 13, False, WD_ALIGN_PARAGRAPH.CENTER, 120)

    add_job1_paragraph(doc, "امضاء المعني", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 8)

    doc.save(output_path)

    try:
        if sys.platform.startswith("win"):
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", output_path])
        else:
            subprocess.Popen(["xdg-open", output_path])
    except Exception:
        pass


master_exam_entries = {}
master_calendar_month = date.today().month
master_calendar_year = date.today().year
master_birth_calendar_month = date.today().month
master_birth_calendar_year = date.today().year


def make_master_entry(x, y, w, placeholder, field_key):
    value = master_exam_entries.get(field_key, "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b8b8b8",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            master_exam_entries[field_key] = ""

    def save_value(event=None):
        val = entry.get()
        master_exam_entries[field_key] = "" if val == placeholder else val

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_master_exam_form():
    global current_page
    current_page = "master_exam_form"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)

    start_y = int(height * 0.055)
    gap = int(height * 0.062)

    right_fields = [
        ("الاسم", "first_name"),
        ("اللقب", "last_name"),
        ("رقم الهاتف", "phone"),
        ("العنوان", "address"),
        ("تاريخ الميلاد", "birth_date"),
        ("تاريخ الطلب", "request_date"),
        ("إلى السيد رئيس قسم", "department_head"),
        ("جامعة", "header_university"),
        ("لولاية", "header_state"),
    ]

    left_fields = [
        ("متابعة دراسة الماستر في", "master_field"),
        ("التخصص", "master_specialty"),
        ("بجامعة", "study_university"),
        ("السنة الجامعية", "academic_year"),
        ("شهادة ليسانس LMD", "license_degree"),
        ("تخصص", "license_specialty"),
        ("من طرف جامعة", "license_university"),
        ("لولاية", "license_state"),
        ("الدفعة", "promotion"),
    ]

    for index, (placeholder, key) in enumerate(right_fields):
        y = start_y + index * gap

        if key == "request_date":
            make_master_entry(right_col_x + 24, y, field_w - 58, placeholder, key)

            cal_icon = canvas.create_text(
                right_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            cal_hitbox = canvas.create_rectangle(
                right_col_x - field_w // 2,
                y - 22,
                right_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_master_cal(event):
                show_master_calendar_picker()

            for item in (cal_icon, cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_master_cal)

        elif key == "birth_date":
            make_master_entry(right_col_x + 24, y, field_w - 58, placeholder, key)

            birth_cal_icon = canvas.create_text(
                right_col_x - field_w // 2 + 28,
                y,
                text="📅",
                fill="#000000",
                font=("Arial", 22, "bold")
            )

            birth_cal_hitbox = canvas.create_rectangle(
                right_col_x - field_w // 2,
                y - 22,
                right_col_x - field_w // 2 + 58,
                y + 25,
                fill="",
                outline=""
            )

            def open_birth_cal(event):
                show_master_birth_calendar_picker()

            for item in (birth_cal_icon, birth_cal_hitbox):
                canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
                canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
                canvas.tag_bind(item, "<Button-1>", open_birth_cal)

        else:
            make_master_entry(right_col_x, y, field_w, placeholder, key)

    for index, (placeholder, key) in enumerate(left_fields):
        make_master_entry(left_col_x, start_y + index * gap, field_w, placeholder, key)

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        create_master_exam_word()

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def draw_master_calendar(kind):
    global current_page
    current_page = "master_birth_calendar_picker" if kind == "birth" else "master_calendar_picker"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    if kind == "birth":
        month = master_birth_calendar_month
        year = master_birth_calendar_year
    else:
        month = master_calendar_month
        year = master_calendar_year

    panel_x1 = int(width * 0.24)
    panel_x2 = int(width * 0.82)
    panel_y1 = int(height * 0.12)
    panel_y2 = int(height * 0.82)

    rounded_home_rect(
        panel_x1,
        panel_y1,
        panel_x2,
        panel_y2,
        r=18,
        fill="#ffffff",
        outline="#dddddd",
        width=2
    )

    canvas.create_text(
        width // 2,
        panel_y1 + 55,
        text=f"{calendar.month_name[month]} {year}",
        fill="#000000",
        font=("Arial", 28, "bold")
    )

    prev_btn = canvas.create_text(panel_x1 + 70, panel_y1 + 55, text="‹", fill="#000000", font=("Arial", 42, "bold"))
    next_btn = canvas.create_text(panel_x2 - 70, panel_y1 + 55, text="›", fill="#000000", font=("Arial", 42, "bold"))

    def prev_month(event):
        global master_calendar_month, master_calendar_year, master_birth_calendar_month, master_birth_calendar_year
        if kind == "birth":
            master_birth_calendar_month -= 1
            if master_birth_calendar_month < 1:
                master_birth_calendar_month = 12
                master_birth_calendar_year -= 1
            show_master_birth_calendar_picker()
        else:
            master_calendar_month -= 1
            if master_calendar_month < 1:
                master_calendar_month = 12
                master_calendar_year -= 1
            show_master_calendar_picker()

    def next_month(event):
        global master_calendar_month, master_calendar_year, master_birth_calendar_month, master_birth_calendar_year
        if kind == "birth":
            master_birth_calendar_month += 1
            if master_birth_calendar_month > 12:
                master_birth_calendar_month = 1
                master_birth_calendar_year += 1
            show_master_birth_calendar_picker()
        else:
            master_calendar_month += 1
            if master_calendar_month > 12:
                master_calendar_month = 1
                master_calendar_year += 1
            show_master_calendar_picker()

    for item, cmd in [(prev_btn, prev_month), (next_btn, next_month)]:
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", cmd)

    days_header = ["Sat", "Sun", "Mon", "Tue", "Wed", "Thu", "Fri"]
    grid_x1 = panel_x1 + 80
    grid_x2 = panel_x2 - 80
    col_w = (grid_x2 - grid_x1) // 7
    start_y = panel_y1 + 120
    row_h = 58

    for i, d in enumerate(days_header):
        canvas.create_text(
            grid_x1 + i * col_w + col_w // 2,
            start_y,
            text=d,
            fill="#555555",
            font=("Arial", 14, "bold")
        )

    cal = calendar.Calendar(firstweekday=5)
    days = list(cal.itermonthdays(year, month))

    for index, day in enumerate(days):
        row = index // 7
        col = index % 7
        x = grid_x1 + col * col_w + col_w // 2
        y = start_y + 45 + row * row_h

        if day == 0:
            continue

        day_box = rounded_home_rect(
            x - 22,
            y - 20,
            x + 22,
            y + 20,
            r=10,
            fill="#ffffff",
            outline="#dddddd",
            width=1
        )

        day_text = canvas.create_text(
            x,
            y,
            text=str(day),
            fill="#000000",
            font=("Arial", 15, "bold")
        )

        def choose_day(event, selected_day=day):
            if kind == "birth":
                master_exam_entries["birth_date"] = f"{selected_day:02d}/{master_birth_calendar_month:02d}/{master_birth_calendar_year}"
            else:
                master_exam_entries["request_date"] = f"{selected_day:02d}/{master_calendar_month:02d}/{master_calendar_year}"
            show_master_exam_form()

        for item in (day_box, day_text):
            canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
            canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
            canvas.tag_bind(item, "<Button-1>", choose_day)


def show_master_calendar_picker():
    draw_master_calendar("request")


def show_master_birth_calendar_picker():
    draw_master_calendar("birth")


def create_master_exam_word():
    output_dir = os.path.join(os.path.expanduser("~"), "IDARA_DZ_Outputs")
    os.makedirs(output_dir, exist_ok=True)

    first_name = master_exam_entries.get("first_name", "").strip()
    last_name = master_exam_entries.get("last_name", "").strip()
    phone = master_exam_entries.get("phone", "").strip()
    address = master_exam_entries.get("address", "").strip()
    birth_date = master_exam_entries.get("birth_date", "").strip()
    request_date = master_exam_entries.get("request_date", "").strip()
    department_head = master_exam_entries.get("department_head", "").strip()
    header_university = master_exam_entries.get("header_university", "").strip()
    header_state = master_exam_entries.get("header_state", "").strip()
    master_field = master_exam_entries.get("master_field", "").strip()
    master_specialty = master_exam_entries.get("master_specialty", "").strip()
    study_university = master_exam_entries.get("study_university", "").strip()
    academic_year = master_exam_entries.get("academic_year", "").strip()
    license_degree = master_exam_entries.get("license_degree", "").strip()
    license_specialty = master_exam_entries.get("license_specialty", "").strip()
    license_university = master_exam_entries.get("license_university", "").strip()
    license_state = master_exam_entries.get("license_state", "").strip()
    promotion = master_exam_entries.get("promotion", "").strip()

    full_name = f"{last_name} {first_name}".strip()
    safe_name = full_name if full_name else "بدون_اسم"
    safe_name = safe_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    output_path = os.path.join(output_dir, f"مسابقة_الماستر_{safe_name}.docx")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.right_margin = Cm(1.5)
    section.left_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(f"التاريخ: {request_date if request_date else '-- / -- / 2026'}.")
    set_job1_run_font(r, 13, True)

    add_job1_paragraph(doc, f"الاسم: {first_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"اللقب: {last_name}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"الهاتف: {phone}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"العنوان: {address}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 2)
    add_job1_paragraph(doc, f"تاريخ الميلاد: {birth_date}", 13, True, WD_ALIGN_PARAGRAPH.RIGHT, 20)

    add_job1_paragraph(doc, f"الى السيد رئيس قسم: {department_head}", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 10)
    add_job1_paragraph(doc, f"جامعة {header_university} لولاية {header_state}", 13, False, WD_ALIGN_PARAGRAPH.LEFT, 62)

    add_job1_paragraph(doc, "الموضوع:", 14, True, WD_ALIGN_PARAGRAPH.RIGHT, 8)
    add_job1_paragraph(doc, "طلب المشاركة في مسابقة الالتحاق بالماستر", 14, True, WD_ALIGN_PARAGRAPH.CENTER, 36)

    body_1 = (
        f"لي عظيم الشرف أن أتقدم إلى سيادتكم بطلبي هذا والمتمثل في طلب تسجيلي ضمن قوائم المترشحين لمتابعة دراسة الماستر "
        f"في {master_field} تخصص {master_specialty} بجامعة {study_university} في إطار السنة الجامعية {academic_year} "
        "بغية مواصلة الدراسة الأكاديمية."
    )

    body_2 = (
        f"علما أني متحصل على شهادة ليسانس نظام * LMD * {license_degree} تخصص {license_specialty} "
        f"من طرف جامعة {license_university} لولاية {license_state} دفعة {promotion}."
    )

    add_job1_paragraph(doc, body_1, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 18)
    add_job1_paragraph(doc, body_2, 13, False, WD_ALIGN_PARAGRAPH.RIGHT, 22)
    add_job1_paragraph(doc, "في انتظار ردكم تقبلوا منا سيدي فائق التقدير والاحترام.", 13, False, WD_ALIGN_PARAGRAPH.CENTER, 120)

    add_job1_paragraph(doc, "امضاء المعني", 13, True, WD_ALIGN_PARAGRAPH.LEFT, 8)

    doc.save(output_path)

    try:
        if sys.platform.startswith("win"):
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", output_path])
        else:
            subprocess.Popen(["xdg-open", output_path])
    except Exception:
        pass


# ============================================================
# Dynamic request-card engine inside "طلب خطي"
# ============================================================

DYNAMIC_DB_DIR = os.path.join(os.path.expanduser("~"), "IDARA_DZ")
DYNAMIC_TEMPLATES_DIR = os.path.join(DYNAMIC_DB_DIR, "templates")
DYNAMIC_OUTPUT_DIR = os.path.join(DYNAMIC_DB_DIR, "outputs")
DYNAMIC_DB_PATH = os.path.join(DYNAMIC_DB_DIR, "idara_dynamic.db")

os.makedirs(DYNAMIC_DB_DIR, exist_ok=True)
os.makedirs(DYNAMIC_TEMPLATES_DIR, exist_ok=True)
os.makedirs(DYNAMIC_OUTPUT_DIR, exist_ok=True)

dynamic_form_values = {}
dynamic_current_card_id = None
dynamic_context_menu_items = []


def init_dynamic_database():
    conn = sqlite3.connect(DYNAMIC_DB_PATH)
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS dynamic_cards (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            template_path TEXT,
            is_visible INTEGER DEFAULT 1,
            sort_order INTEGER DEFAULT 999,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS dynamic_fields (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            card_id INTEGER NOT NULL,
            field_name TEXT NOT NULL,
            field_type TEXT DEFAULT 'text',
            sort_order INTEGER DEFAULT 999,
            FOREIGN KEY(card_id) REFERENCES dynamic_cards(id) ON DELETE CASCADE
        )
    """)

    conn.commit()
    conn.close()


def db_fetchall(query, params=()):
    init_dynamic_database()
    conn = sqlite3.connect(DYNAMIC_DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(query, params)
    rows = cur.fetchall()
    conn.close()
    return rows


def db_execute(query, params=()):
    init_dynamic_database()
    conn = sqlite3.connect(DYNAMIC_DB_PATH)
    cur = conn.cursor()
    cur.execute(query, params)
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id


def get_dynamic_cards():
    return db_fetchall("SELECT * FROM dynamic_cards WHERE is_visible=1 ORDER BY sort_order ASC, id ASC")


def get_dynamic_fields(card_id):
    return db_fetchall("SELECT * FROM dynamic_fields WHERE card_id=? ORDER BY sort_order ASC, id ASC", (card_id,))


def save_dynamic_fields(card_id, fields_text):
    for idx, line in enumerate((fields_text or "").splitlines()):
        line = line.strip()
        if not line:
            continue

        if "|" in line:
            name, field_type = line.split("|", 1)
            name = name.strip()
            field_type = field_type.strip()
        else:
            name = line.strip()
            field_type = "text"

        if name:
            db_execute(
                "INSERT INTO dynamic_fields(card_id, field_name, field_type, sort_order) VALUES (?, ?, ?, ?)",
                (card_id, name, field_type, idx)
            )


def add_dynamic_card(title, fields_text, template_path=""):
    title = (title or "").strip()
    if not title:
        return None

    stored_template = ""
    if template_path and os.path.exists(template_path):
        ext = os.path.splitext(template_path)[1].lower()
        stored_template = os.path.join(DYNAMIC_TEMPLATES_DIR, f"{uuid.uuid4().hex}{ext}")
        shutil.copy2(template_path, stored_template)

    rows = db_fetchall("SELECT MAX(sort_order) AS max_order FROM dynamic_cards")
    sort_order = 999
    if rows and rows[0]["max_order"] is not None:
        sort_order = int(rows[0]["max_order"]) + 1

    card_id = db_execute(
        "INSERT INTO dynamic_cards(title, template_path, is_visible, sort_order, created_at) VALUES (?, ?, 1, ?, ?)",
        (title, stored_template, sort_order, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )

    save_dynamic_fields(card_id, fields_text)
    return card_id


def update_dynamic_card(card_id, title, fields_text, template_path=""):
    title = (title or "").strip()
    if not title:
        return

    if template_path and os.path.exists(template_path):
        ext = os.path.splitext(template_path)[1].lower()
        stored_template = os.path.join(DYNAMIC_TEMPLATES_DIR, f"{uuid.uuid4().hex}{ext}")
        shutil.copy2(template_path, stored_template)
        db_execute("UPDATE dynamic_cards SET title=?, template_path=? WHERE id=?", (title, stored_template, card_id))
    else:
        db_execute("UPDATE dynamic_cards SET title=? WHERE id=?", (title, card_id))

    db_execute("DELETE FROM dynamic_fields WHERE card_id=?", (card_id,))
    save_dynamic_fields(card_id, fields_text)


def delete_dynamic_card(card_id):
    db_execute("DELETE FROM dynamic_fields WHERE card_id=?", (card_id,))
    db_execute("DELETE FROM dynamic_cards WHERE id=?", (card_id,))


def replace_text_in_docx(doc, replacements):
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        changed = False

        for key, value in replacements.items():
            marker = "{{" + key + "}}"
            if marker in full_text:
                full_text = full_text.replace(marker, str(value))
                changed = True

        if changed:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                paragraph.add_run(full_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    changed = False
                    for key, value in replacements.items():
                        marker = "{{" + key + "}}"
                        if marker in full_text:
                            full_text = full_text.replace(marker, str(value))
                            changed = True
                    if changed:
                        for run in paragraph.runs:
                            run.text = ""
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text
                        else:
                            paragraph.add_run(full_text)


def open_dynamic_generated_word(card_id):
    rows = db_fetchall("SELECT * FROM dynamic_cards WHERE id=?", (card_id,))
    if not rows:
        return

    card = rows[0]
    template_path = card["template_path"]

    if not template_path or not os.path.exists(template_path):
        show_dynamic_message("لا يوجد نموذج Word مربوط بهذه البطاقة.")
        return

    doc = Document(template_path)
    replace_text_in_docx(doc, dynamic_form_values)

    safe_title = card["title"].replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    client_name = (dynamic_form_values.get("اللقب", "") + "_" + dynamic_form_values.get("الاسم", "")).strip("_") or "بدون_اسم"
    client_name = client_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")

    output_path = os.path.join(
        DYNAMIC_OUTPUT_DIR,
        f"{safe_title}_{client_name}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
    )

    doc.save(output_path)

    try:
        if sys.platform.startswith("win"):
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", output_path])
        else:
            subprocess.Popen(["xdg-open", output_path])
    except Exception:
        pass


def show_dynamic_message(message):
    clear_screen()
    width = root.winfo_width()
    height = root.winfo_height()
    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")
    center_x = 120 + (width - 120) // 2
    canvas.create_text(center_x, height // 2, text=message, fill="#000000", font=("Arial", 26, "bold"))


def show_settings_main():
    show_settings_main()


def show_settings_main():
    global current_page
    current_page = "settings"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("settings")


def make_dynamic_entry(x, y, w, placeholder, field_name, field_type):
    value = dynamic_form_values.get(field_name, "")

    canvas.create_line(x - w // 2, y + 18, x + w // 2, y + 18, fill="#b8b8b8", width=2)

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            dynamic_form_values[field_name] = ""

    def save_value(event=None):
        val = entry.get()
        dynamic_form_values[field_name] = "" if val == placeholder else val

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_dynamic_card_form(card_id):
    global current_page, dynamic_current_card_id, dynamic_form_values
    current_page = "dynamic_card_form"
    dynamic_current_card_id = card_id
    dynamic_form_values = {}

    rows = db_fetchall("SELECT * FROM dynamic_cards WHERE id=?", (card_id,))
    if not rows:
        return

    card = rows[0]
    fields = get_dynamic_fields(card_id)

    clear_screen()
    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    center_x = 120 + (width - 120) // 2
    canvas.create_text(center_x, 55, text=card["title"], fill="#000000", font=("Arial", 34, "bold"))

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)
    start_y = int(height * 0.12)
    gap = int(height * 0.075)

    for idx, field in enumerate(fields):
        col_x = right_col_x if idx % 2 == 0 else left_col_x
        y = start_y + (idx // 2) * gap
        make_dynamic_entry(col_x, y, field_w, field["field_name"], field["field_name"], field["field_type"])

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        open_dynamic_generated_word(card_id)

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def show_dynamic_card_editor(card_id=None):
    global current_page
    current_page = "dynamic_card_editor"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    center_x = 120 + (width - 120) // 2
    editing = card_id is not None

    canvas.create_text(
        center_x,
        70,
        text="تعديل بطاقة طلب خطي" if editing else "إضافة بطاقة طلب خطي",
        fill="#000000",
        font=("Arial", 34, "bold")
    )

    old_title = ""
    old_fields = ""
    old_template = ""

    if editing:
        rows = db_fetchall("SELECT * FROM dynamic_cards WHERE id=?", (card_id,))
        if rows:
            old_title = rows[0]["title"]
            old_template = rows[0]["template_path"] or ""

        fields = get_dynamic_fields(card_id)
        old_fields = "\n".join([f'{f["field_name"]} | {f["field_type"]}' for f in fields])

    canvas.create_text(center_x + 300, 145, text="اسم البطاقة", fill="#000000", font=("Arial", 17, "bold"), anchor="e")

    title_entry = tk.Entry(root, font=("Arial", 16, "bold"), justify="right", bd=1, relief="solid")
    title_entry.insert(0, old_title)
    canvas.create_window(center_x, 185, window=title_entry, width=600, height=44)

    canvas.create_text(center_x + 300, 250, text="خانات الاستمارة", fill="#000000", font=("Arial", 17, "bold"), anchor="e")
    canvas.create_text(center_x, 282, text="كل خانة في سطر. مثال: الاسم | text     تاريخ الطلب | date", fill="#777777", font=("Arial", 12, "bold"))

    fields_text = tk.Text(root, font=("Arial", 15), bd=1, relief="solid", wrap="word")
    fields_text.insert("1.0", old_fields)
    canvas.create_window(center_x, 390, window=fields_text, width=600, height=165)

    template_path_value = {"path": ""}

    template_label = canvas.create_text(
        center_x,
        515,
        text=os.path.basename(old_template) if old_template else "لم يتم اختيار نموذج Word",
        fill="#000000" if old_template else "#777777",
        font=("Arial", 13, "bold")
    )

    choose_btn = rounded_home_rect(center_x - 310, 545, center_x - 75, 595, r=12, fill="#f4f4f4", outline="#d5d5d5", width=1)
    choose_text = canvas.create_text(center_x - 192, 570, text="اختيار / تغيير نموذج Word", fill="#000000", font=("Arial", 13, "bold"))

    def choose_template(event):
        path = filedialog.askopenfilename(title="اختر نموذج Word", filetypes=[("Word files", "*.docx")])
        if path:
            template_path_value["path"] = path
            canvas.itemconfig(template_label, text=os.path.basename(path), fill="#000000")

    for item in (choose_btn, choose_text):
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", choose_template)

    save_btn = rounded_home_rect(center_x + 75, 545, center_x + 310, 595, r=12, fill="#000000", outline="#000000", width=1)
    save_text = canvas.create_text(center_x + 192, 570, text="حفظ", fill="#ffffff", font=("Arial", 15, "bold"))

    def save_card(event):
        title = title_entry.get().strip()
        fields = fields_text.get("1.0", "end").strip()

        if editing:
            update_dynamic_card(card_id, title, fields, template_path_value["path"])
        else:
            add_dynamic_card(title, fields, template_path_value["path"])

        show_written_request_menu()

    for item in (save_btn, save_text):
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", save_card)


def clear_dynamic_context_menu():
    global dynamic_context_menu_items
    for item in dynamic_context_menu_items:
        try:
            canvas.delete(item)
        except:
            pass
    dynamic_context_menu_items = []


def show_card_context_menu(x, y, title):
    global dynamic_context_menu_items
    clear_dynamic_context_menu()

    rows = db_fetchall("SELECT * FROM dynamic_cards WHERE title=? AND is_visible=1", (title,))
    if not rows:
        return

    card_id = rows[0]["id"]

    box = rounded_home_rect(x, y, x + 155, y + 92, r=10, fill="#ffffff", outline="#d0d0d0", width=1)
    edit_text = canvas.create_text(x + 78, y + 28, text="✏️ تعديل", fill="#000000", font=("Arial", 14, "bold"))
    delete_text = canvas.create_text(x + 78, y + 66, text="🗑️ حذف", fill="#d62323", font=("Arial", 14, "bold"))

    dynamic_context_menu_items.extend([box, edit_text, delete_text])

    def edit_click(event):
        clear_dynamic_context_menu()
        show_dynamic_card_editor(card_id)

    def delete_click(event):
        clear_dynamic_context_menu()
        delete_dynamic_card(card_id)
        show_written_request_menu()

    for item in (edit_text,):
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", edit_click)

    for item in (delete_text,):
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", delete_click)

def make_light_underline_entry(x, y, w, placeholder, field_key):
    value = job_form_entries.get(field_key, "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b8b8b8",
        width=2
    )

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            job_form_entries[field_key] = ""
            root.after(200, clear_suggestions)

    def save_value(event=None):
        val = entry.get()
        if val == placeholder:
            job_form_entries[field_key] = ""
            clear_suggestions()
        else:
            job_form_entries[field_key] = val
            draw_client_suggestions(x, y, w, field_key, val)

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def draw_light_request_type_field(x, y, w):
    selected_value = job_form_entries.get("request_type", "")

    canvas.create_line(
        x - w // 2,
        y + 18,
        x + w // 2,
        y + 18,
        fill="#b8b8b8",
        width=2
    )

    arrow = canvas.create_text(
        x - w // 2 + 18,
        y,
        text="▼",
        fill="#b5b5b5",
        font=("Arial", 18, "bold")
    )

    text = canvas.create_text(
        x + w // 2,
        y,
        text=selected_value if selected_value else "نوع الطلب",
        fill="#111111" if selected_value else "#b5b5b5",
        font=("Arial", 16, "bold"),
        anchor="e"
    )

    hitbox = canvas.create_rectangle(
        x - w // 2,
        y - 18,
        x + w // 2,
        y + 25,
        fill="",
        outline=""
    )

    def enter(event):
        root.config(cursor="hand2")

    def leave(event):
        root.config(cursor="")

    def click(event):
        open_request_type_dropdown()

    for item in (arrow, text, hitbox):
        canvas.tag_bind(item, "<Enter>", enter)
        canvas.tag_bind(item, "<Leave>", leave)
        canvas.tag_bind(item, "<Button-1>", click)

    if request_type_dropdown_open:
        draw_request_type_dropdown(x, y + 28, w)


def show_job_request_form():
    global current_page
    current_page = "job_request_form"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    # White form page exactly like the requested screenshot
    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    sidebar_w = 120
    content_x1 = sidebar_w
    content_w = width - sidebar_w
    center_x = content_x1 + content_w // 2

    # Column positions and field geometry
    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)

    start_y = int(height * 0.10)
    gap = int(height * 0.112)

    # Right column
    right_fields = [
        ("الاسم", "first_name"),
        ("اللقب", "last_name"),
        ("تاريخ و مكان الميلاد", "birth_info"),
        ("العنوان الكامل", "address"),
        ("رقم بطاقة التعريف", "id_card"),
        ("رقم الهاتف", "phone"),
    ]

    # Left column
    left_fields = [
        ("السيد/ الجهة المستقبلة", "recipient"),
        ("المنصب", "position"),
        ("الشهادة", "degree"),
        ("التخصص", "specialty"),
        ("تاريخ الطلب", "date"),
        ("نوع الطلب", "request_type"),
    ]

    for index, (placeholder, key) in enumerate(right_fields):
        make_light_underline_entry(
            right_col_x,
            start_y + index * gap,
            field_w,
            placeholder,
            key
        )

    for index, (placeholder, key) in enumerate(left_fields):
        y = start_y + index * gap

        if key == "request_type":
            draw_light_request_type_field(left_col_x, y, field_w)
        else:
            make_light_underline_entry(
                left_col_x,
                y,
                field_w,
                placeholder,
                key
            )

    # Calendar icon next to date, like the screenshot
    date_y = start_y + 4 * gap
    cal_icon = canvas.create_text(
        left_col_x - field_w // 2 + 24,
        date_y,
        text="▣",
        fill="#000000",
        font=("Arial", 28, "bold")
    )

    cal_hitbox = canvas.create_rectangle(
        left_col_x - field_w // 2,
        date_y - 22,
        left_col_x - field_w // 2 + 58,
        date_y + 25,
        fill="",
        outline=""
    )

    def open_cal(event):
        show_calendar_picker()

    for item in (cal_icon, cal_hitbox):
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", open_cal)

    # Preview text button bottom right
    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        show_job_request_preview()

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def scroll_job_form(event):
    global job_form_scroll

    if current_page != "job_request_form":
        return

    if event.delta < 0:
        job_form_scroll += 1
    else:
        job_form_scroll -= 1

    job_form_scroll = max(0, min(job_form_scroll, 6))
    show_job_request_form()


def show_job_request_preview():
    selected_type = job_form_entries.get("request_type", "").strip()

    if selected_type and selected_type != "طلب توظيف (عام)":
        try:
            messagebox.showinfo(
                "تنبيه",
                "حاليًا تم تجهيز نموذج Word لطلب توظيف (عام) فقط.\nسنضيف بقية النماذج واحدًا واحدًا لاحقًا."
            )
        except:
            pass

    try:
        output_path = create_job_request_word()
        archive_path = archive_generated_file(output_path)
        save_or_update_client_from_form(archive_path)
        open_generated_file(output_path)
    except Exception as e:
        try:
            messagebox.showerror("خطأ", f"تعذر إنشاء ملف Word:\n{e}")
        except:
            pass


def show_section(section):
    global current_page
    current_page = section
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    if width < 10 or height < 10:
        width, height = 1280, 720

    draw_image(resource_path("assets/background.jpg"), width, height)

    titles = {
        "electronic": "خدمات إلكترونية",
        "archive": "أرشيف",
    }

    canvas.create_rectangle(
        int(width * 0.18),
        int(height * 0.22),
        int(width * 0.82),
        int(height * 0.72),
        fill="#173b38",
        outline="#d7c28a",
        width=2
    )

    canvas.create_text(
        width // 2,
        int(height * 0.34),
        text=titles.get(section, "صفحة"),
        fill="#f4f4f4",
        font=("Arial", 44, "bold")
    )

    canvas.create_text(
        width // 2,
        int(height * 0.50),
        text="سنقوم ببناء هذا القسم لاحقًا.",
        fill="#e8e1d5",
        font=("Arial", 24, "bold")
    )

    draw_back_button(show_home)


def show_about():
    global current_page
    current_page = "about"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("about")

    center_x = 120 + (width - 120) // 2

    canvas.create_text(center_x, int(height * 0.20), text="حول البرنامج", fill="#000000", font=("Arial", 44, "bold"))

    text = (
        "IDARA DZ\n\n"
        "برنامج مكتبي مخصص لتنظيم خدمات المكتبة والخدمات الإدارية.\n"
        "يساعد على تسيير الوثائق، الخدمات الإلكترونية، والأرشيف بطريقة سهلة ومنظمة.\n\n"
        "الإصدار: 1.0.0\n"
        "© 2026"
    )

    canvas.create_text(
        center_x,
        int(height * 0.48),
        text=text,
        fill="#555555",
        font=("Arial", 22, "bold"),
        justify="center",
        width=int((width - 120) * 0.70)
    )


def on_resize(event):
    if event.widget == root:
        if current_page == "home":
            show_home()
        elif current_page == "about":
            show_about()
        elif current_page == "settings":
            show_settings_main()
        elif current_page == "documents":
            show_documents()
        elif current_page == "written_request":
            show_written_request()
        elif current_page == "written_request_template":
            show_written_request()
        elif current_page == "job_request_form":
            show_job_request_form()
        elif current_page == "job_request_1_form":
            show_job_request_1_form()
        elif current_page == "job_request_2_form":
            show_job_request_2_form()
        elif current_page == "job2_calendar_picker":
            show_job2_calendar_picker()
        elif current_page == "customs_exam_form":
            show_customs_exam_form()
        elif current_page == "customs_calendar_picker":
            show_customs_calendar_picker()
        elif current_page == "customs_birth_calendar_picker":
            show_customs_birth_calendar_picker()
        elif current_page == "police_exam_form":
            show_police_exam_form()
        elif current_page == "police_calendar_picker":
            show_police_calendar_picker()
        elif current_page == "police_birth_calendar_picker":
            show_police_birth_calendar_picker()
        elif current_page == "civil_protection_exam_form":
            show_civil_protection_exam_form()
        elif current_page == "civil_calendar_picker":
            show_civil_calendar_picker()
        elif current_page == "civil_birth_calendar_picker":
            show_civil_birth_calendar_picker()
        elif current_page == "pre_employment_contract_form":
            show_pre_employment_contract_form()
        elif current_page == "pre_employment_calendar_picker":
            show_pre_employment_calendar_picker()
        elif current_page == "master_exam_form":
            show_master_exam_form()
        elif current_page == "master_calendar_picker":
            show_master_calendar_picker()
        elif current_page == "master_birth_calendar_picker":
            show_master_birth_calendar_picker()
        elif current_page == "dynamic_card_editor":
            show_dynamic_card_editor()
        elif current_page == "dynamic_card_form":
            show_dynamic_card_form(dynamic_current_card_id)
        elif current_page == "settings":
            show_settings_main()
        elif current_page == "job1_calendar_picker":
            show_job1_calendar_picker()
        elif current_page == "written_request_menu":
            show_written_request_menu()
        elif current_page == "job_request_preview":
            show_job_request_preview()
        elif current_page == "calendar_picker":
            show_calendar_picker()
        elif current_page in ["honor_statement", "cv", "invoice"]:
            show_document_type(current_page)
        else:
            show_section(current_page)



def on_request_type_key(event):
    if current_page != "job_request_form" or not request_type_dropdown_open:
        return

    if event.keysym == "Down":
        move_request_type_selection(1)
    elif event.keysym == "Up":
        move_request_type_selection(-1)
    elif event.keysym == "Return":
        confirm_request_type_selection()
    elif event.keysym == "Escape":
        close_request_type_dropdown()




# ============================================================
# FIXED DYNAMIC CARDS SYSTEM - stable Arabic workflow
# ============================================================

HIDDEN_BUILTIN_CARDS_FILE = os.path.join(os.path.expanduser("~"), "IDARA_DZ", "hidden_builtin_cards.json")
_long_press_after_id = None
_long_press_title = None


def load_hidden_builtin_cards():
    try:
        if os.path.exists(HIDDEN_BUILTIN_CARDS_FILE):
            with open(HIDDEN_BUILTIN_CARDS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            return set(data if isinstance(data, list) else [])
    except Exception:
        pass
    return set()


def save_hidden_builtin_cards(hidden_set):
    os.makedirs(os.path.dirname(HIDDEN_BUILTIN_CARDS_FILE), exist_ok=True)
    with open(HIDDEN_BUILTIN_CARDS_FILE, "w", encoding="utf-8") as f:
        json.dump(list(hidden_set), f, ensure_ascii=False, indent=2)


def hide_builtin_card(title):
    hidden = load_hidden_builtin_cards()
    hidden.add(title)
    save_hidden_builtin_cards(hidden)


def init_dynamic_database():
    conn = sqlite3.connect(DYNAMIC_DB_PATH)
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS dynamic_cards (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            template_path TEXT,
            is_visible INTEGER DEFAULT 1,
            sort_order INTEGER DEFAULT 999,
            created_at TEXT
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS dynamic_fields (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            card_id INTEGER NOT NULL,
            field_name TEXT NOT NULL,
            field_type TEXT DEFAULT 'text',
            sort_order INTEGER DEFAULT 999,
            FOREIGN KEY(card_id) REFERENCES dynamic_cards(id) ON DELETE CASCADE
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS dynamic_drafts (
            card_id INTEGER NOT NULL,
            field_name TEXT NOT NULL,
            field_value TEXT,
            PRIMARY KEY(card_id, field_name)
        )
    """)

    conn.commit()
    conn.close()


def db_fetchall(query, params=()):
    init_dynamic_database()
    conn = sqlite3.connect(DYNAMIC_DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(query, params)
    rows = cur.fetchall()
    conn.close()
    return rows


def db_execute(query, params=()):
    init_dynamic_database()
    conn = sqlite3.connect(DYNAMIC_DB_PATH)
    cur = conn.cursor()
    cur.execute(query, params)
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id


def infer_field_type(field_name):
    name = (field_name or "").strip()
    if "تاريخ" in name or "ميلاد" in name:
        return "date"
    return "text"


def normalize_fields_text(fields_text):
    lines = []
    for line in (fields_text or "").splitlines():
        clean = line.strip()
        if not clean:
            continue
        if "|" in clean:
            name = clean.split("|", 1)[0].strip()
        else:
            name = clean
        if name:
            lines.append(name)
    return lines


def save_dynamic_fields(card_id, fields_text):
    fields = normalize_fields_text(fields_text)
    for idx, name in enumerate(fields):
        db_execute(
            "INSERT INTO dynamic_fields(card_id, field_name, field_type, sort_order) VALUES (?, ?, ?, ?)",
            (card_id, name, infer_field_type(name), idx)
        )


def add_dynamic_card(title, fields_text, template_path=""):
    title = (title or "").strip()
    if not title:
        messagebox.showwarning("تنبيه", "اكتب اسم البطاقة أولا.")
        return None

    stored_template = ""
    if template_path and os.path.exists(template_path):
        ext = os.path.splitext(template_path)[1].lower()
        stored_template = os.path.join(DYNAMIC_TEMPLATES_DIR, f"{uuid.uuid4().hex}{ext}")
        shutil.copy2(template_path, stored_template)

    rows = db_fetchall("SELECT MAX(sort_order) AS max_order FROM dynamic_cards")
    sort_order = 999 if not rows or rows[0]["max_order"] is None else int(rows[0]["max_order"]) + 1

    card_id = db_execute(
        "INSERT INTO dynamic_cards(title, template_path, is_visible, sort_order, created_at) VALUES (?, ?, 1, ?, ?)",
        (title, stored_template, sort_order, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )

    save_dynamic_fields(card_id, fields_text)
    return card_id


def update_dynamic_card(card_id, title, fields_text, template_path=""):
    title = (title or "").strip()
    if not title:
        messagebox.showwarning("تنبيه", "اكتب اسم البطاقة أولا.")
        return

    if template_path and os.path.exists(template_path):
        ext = os.path.splitext(template_path)[1].lower()
        stored_template = os.path.join(DYNAMIC_TEMPLATES_DIR, f"{uuid.uuid4().hex}{ext}")
        shutil.copy2(template_path, stored_template)
        db_execute("UPDATE dynamic_cards SET title=?, template_path=? WHERE id=?", (title, stored_template, card_id))
    else:
        db_execute("UPDATE dynamic_cards SET title=? WHERE id=?", (title, card_id))

    db_execute("DELETE FROM dynamic_fields WHERE card_id=?", (card_id,))
    save_dynamic_fields(card_id, fields_text)


def delete_dynamic_card(card_id):
    db_execute("DELETE FROM dynamic_drafts WHERE card_id=?", (card_id,))
    db_execute("DELETE FROM dynamic_fields WHERE card_id=?", (card_id,))
    db_execute("DELETE FROM dynamic_cards WHERE id=?", (card_id,))


def load_dynamic_draft(card_id):
    rows = db_fetchall("SELECT field_name, field_value FROM dynamic_drafts WHERE card_id=?", (card_id,))
    return {r["field_name"]: r["field_value"] or "" for r in rows}


def save_dynamic_draft_value(card_id, field_name, value):
    db_execute(
        "INSERT OR REPLACE INTO dynamic_drafts(card_id, field_name, field_value) VALUES (?, ?, ?)",
        (card_id, field_name, value)
    )


def clear_dynamic_context_menu():
    global dynamic_context_menu_items
    for item in dynamic_context_menu_items:
        try:
            canvas.delete(item)
        except Exception:
            pass
    dynamic_context_menu_items = []


def show_card_context_menu(x, y, title):
    global dynamic_context_menu_items
    clear_dynamic_context_menu()

    dynamic_rows = db_fetchall("SELECT * FROM dynamic_cards WHERE title=? AND is_visible=1", (title,))
    is_dynamic = bool(dynamic_rows)
    card_id = dynamic_rows[0]["id"] if is_dynamic else None

    box = rounded_home_rect(x, y, x + 155, y + 92, r=10, fill="#ffffff", outline="#d0d0d0", width=1)
    edit_text = canvas.create_text(x + 78, y + 28, text="✏️ تعديل", fill="#000000", font=("Arial", 14, "bold"))
    delete_text = canvas.create_text(x + 78, y + 66, text="🗑️ حذف", fill="#d62323", font=("Arial", 14, "bold"))
    dynamic_context_menu_items.extend([box, edit_text, delete_text])

    def edit_click(event):
        clear_dynamic_context_menu()
        if is_dynamic:
            show_dynamic_card_editor(card_id)
        else:
            messagebox.showinfo("تنبيه", "التعديل الكامل متاح للبطاقات التي تضيفها من زر + فقط. البطاقات الأساسية سننقلها لاحقًا للنظام الديناميكي.")

    def delete_click(event):
        clear_dynamic_context_menu()
        if is_dynamic:
            delete_dynamic_card(card_id)
        else:
            hide_builtin_card(title)
        show_written_request_menu()

    for item in (edit_text, delete_text):
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))

    canvas.tag_bind(edit_text, "<Button-1>", edit_click)
    canvas.tag_bind(delete_text, "<Button-1>", delete_click)


def start_card_long_press(event, title):
    global _long_press_after_id, _long_press_title
    _long_press_title = title

    if _long_press_after_id:
        try:
            root.after_cancel(_long_press_after_id)
        except Exception:
            pass

    _long_press_after_id = root.after(650, lambda: show_card_context_menu(event.x, event.y, title))


def cancel_card_long_press(event=None):
    global _long_press_after_id, _long_press_title
    if _long_press_after_id:
        try:
            root.after_cancel(_long_press_after_id)
        except Exception:
            pass
    _long_press_after_id = None
    _long_press_title = None


def show_dynamic_card_editor(card_id=None):
    global current_page
    current_page = "dynamic_card_editor"
    clear_screen()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    center_x = 120 + (width - 120) // 2
    editing = card_id is not None

    canvas.create_text(
        center_x,
        58,
        text="تعديل بطاقة طلب خطي" if editing else "إضافة بطاقة طلب خطي",
        fill="#000000",
        font=("Arial", 32, "bold")
    )

    old_title = ""
    old_fields = ""
    old_template = ""

    if editing:
        rows = db_fetchall("SELECT * FROM dynamic_cards WHERE id=?", (card_id,))
        if rows:
            old_title = rows[0]["title"]
            old_template = rows[0]["template_path"] or ""
        fields = get_dynamic_fields(card_id)
        old_fields = "\n".join([f["field_name"] for f in fields])

    canvas.create_text(center_x + 300, 128, text="اسم البطاقة", fill="#000000", font=("Arial", 17, "bold"), anchor="e")

    title_entry = tk.Entry(root, font=("Arial", 16, "bold"), justify="right", bd=1, relief="solid")
    title_entry.insert(0, old_title)
    canvas.create_window(center_x, 168, window=title_entry, width=600, height=44)

    canvas.create_text(center_x + 300, 228, text="خانات الاستمارة", fill="#000000", font=("Arial", 17, "bold"), anchor="e")
    canvas.create_text(center_x, 258, text="اكتب أسماء الخانات فقط، كل خانة في سطر. البرنامج يعرف التاريخ تلقائيًا من كلمة تاريخ.", fill="#777777", font=("Arial", 12, "bold"))

    fields_text = tk.Text(root, font=("Arial", 15), bd=1, relief="solid", wrap="word")
    fields_text.tag_configure("right", justify="right")
    fields_text.insert("1.0", old_fields)
    fields_text.tag_add("right", "1.0", "end")
    fields_text.bind("<KeyRelease>", lambda e: fields_text.tag_add("right", "1.0", "end"))
    canvas.create_window(center_x, 375, window=fields_text, width=600, height=170)

    template_path_value = {"path": ""}

    template_label = canvas.create_text(
        center_x,
        502,
        text=os.path.basename(old_template) if old_template else "لم يتم اختيار نموذج Word",
        fill="#000000" if old_template else "#777777",
        font=("Arial", 13, "bold")
    )

    choose_btn = rounded_home_rect(center_x - 310, 535, center_x - 75, 585, r=12, fill="#f4f4f4", outline="#d5d5d5", width=1)
    choose_text = canvas.create_text(center_x - 192, 560, text="اختيار / تغيير نموذج Word", fill="#000000", font=("Arial", 13, "bold"))

    def choose_template(event):
        path = filedialog.askopenfilename(title="اختر نموذج Word", filetypes=[("Word files", "*.docx")])
        if path:
            template_path_value["path"] = path
            canvas.itemconfig(template_label, text=os.path.basename(path), fill="#000000")

    for item in (choose_btn, choose_text):
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", choose_template)

    save_btn = rounded_home_rect(center_x + 75, 535, center_x + 310, 585, r=12, fill="#000000", outline="#000000", width=1)
    save_text = canvas.create_text(center_x + 192, 560, text="حفظ", fill="#ffffff", font=("Arial", 15, "bold"))

    def save_card(event):
        title = title_entry.get().strip()
        fields = fields_text.get("1.0", "end").strip()
        if editing:
            update_dynamic_card(card_id, title, fields, template_path_value["path"])
        else:
            add_dynamic_card(title, fields, template_path_value["path"])
        written_request_search_query = ""
        show_written_request_menu()

    for item in (save_btn, save_text):
        canvas.tag_bind(item, "<Enter>", lambda e: root.config(cursor="hand2"))
        canvas.tag_bind(item, "<Leave>", lambda e: root.config(cursor=""))
        canvas.tag_bind(item, "<Button-1>", save_card)


def make_dynamic_entry(x, y, w, placeholder, field_name, field_type):
    value = dynamic_form_values.get(field_name, "")

    canvas.create_line(x - w // 2, y + 18, x + w // 2, y + 18, fill="#b8b8b8", width=2)

    entry = tk.Entry(
        root,
        font=("Arial", 16, "bold"),
        justify="right",
        bd=0,
        bg="#ffffff",
        fg="#111111",
        insertbackground="#111111"
    )

    if value:
        entry.insert(0, value)
        entry.config(fg="#111111")
    else:
        entry.insert(0, placeholder)
        entry.config(fg="#b5b5b5")

    def focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, "end")
            entry.config(fg="#111111")

    def focus_out(event):
        if not entry.get().strip():
            entry.delete(0, "end")
            entry.insert(0, placeholder)
            entry.config(fg="#b5b5b5")
            dynamic_form_values[field_name] = ""
            if dynamic_current_card_id:
                save_dynamic_draft_value(dynamic_current_card_id, field_name, "")

    def save_value(event=None):
        val = entry.get()
        real_value = "" if val == placeholder else val
        dynamic_form_values[field_name] = real_value
        if dynamic_current_card_id:
            save_dynamic_draft_value(dynamic_current_card_id, field_name, real_value)

    entry.bind("<FocusIn>", focus_in)
    entry.bind("<FocusOut>", focus_out)
    entry.bind("<KeyRelease>", save_value)

    canvas.create_window(x, y, window=entry, width=w, height=34)
    return entry


def show_dynamic_card_form(card_id):
    global current_page, dynamic_current_card_id, dynamic_form_values
    current_page = "dynamic_card_form"
    dynamic_current_card_id = card_id

    rows = db_fetchall("SELECT * FROM dynamic_cards WHERE id=?", (card_id,))
    if not rows:
        return

    dynamic_form_values = load_dynamic_draft(card_id)
    card = rows[0]
    fields = get_dynamic_fields(card_id)

    clear_screen()
    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#ffffff", outline="#ffffff")
    draw_home_sidebar("home")

    center_x = 120 + (width - 120) // 2
    canvas.create_text(center_x, 55, text=card["title"], fill="#000000", font=("Arial", 34, "bold"))

    right_col_x = int(width * 0.80)
    left_col_x = int(width * 0.34)
    field_w = int(width * 0.29)
    start_y = int(height * 0.12)
    gap = int(height * 0.075)

    for idx, field in enumerate(fields):
        col_x = right_col_x if idx % 2 == 0 else left_col_x
        y = start_y + (idx // 2) * gap
        make_dynamic_entry(col_x, y, field_w, field["field_name"], field["field_name"], field["field_type"])

    preview_text = canvas.create_text(
        int(width * 0.90),
        int(height * 0.88),
        text="معاينة",
        fill="#55bfff",
        font=("Arial", 27, "bold"),
        anchor="center"
    )

    def preview_enter(event):
        canvas.itemconfig(preview_text, fill="#1d9fee")
        root.config(cursor="hand2")

    def preview_leave(event):
        canvas.itemconfig(preview_text, fill="#55bfff")
        root.config(cursor="")

    def preview_click(event):
        open_dynamic_generated_word(card_id)

    canvas.tag_bind(preview_text, "<Enter>", preview_enter)
    canvas.tag_bind(preview_text, "<Leave>", preview_leave)
    canvas.tag_bind(preview_text, "<Button-1>", preview_click)


def show_written_request_menu():
    global current_page, written_request_menu_scroll, written_request_search_query
    current_page = "written_request_menu"
    clear_screen()
    init_dynamic_database()

    width = root.winfo_width()
    height = root.winfo_height()

    canvas.create_rectangle(0, 0, width, height, fill="#efefef", outline="#efefef")
    draw_home_sidebar("home")

    search_x = width - 290
    search_w = 280
    canvas.create_line(search_x - search_w, 104, search_x + 10, 104, fill="#bdbdbd", width=2)

    search_entry = tk.Entry(root, bd=0, bg="#efefef", fg="#111111", font=("Arial", 17, "bold"), justify="right", insertbackground="#111111")

    if written_request_search_query:
        search_entry.insert(0, written_request_search_query)
        search_entry.config(fg="#111111")
    else:
        search_entry.insert(0, "بحث")
        search_entry.config(fg="#9b9b9b")

    def s_in(e):
        if search_entry.get() == "بحث":
            search_entry.delete(0, "end")
            search_entry.config(fg="#111111")

    def s_out(e):
        global written_request_search_query
        if not search_entry.get():
            written_request_search_query = ""
            search_entry.insert(0, "بحث")
            search_entry.config(fg="#9b9b9b")

    def do_search(e=None):
        global written_request_search_query, written_request_menu_scroll
        text = search_entry.get().strip()
        if text == "بحث":
            text = ""
        written_request_search_query = text
        written_request_menu_scroll = 0
        show_written_request_menu()

    search_entry.bind("<FocusIn>", s_in)
    search_entry.bind("<FocusOut>", s_out)
    search_entry.bind("<KeyRelease>", do_search)

    canvas.create_window(search_x - 125, 82, window=search_entry, width=search_w - 42, height=30)
    canvas.create_text(search_x + 5, 82, text="⌕", fill="#a5a5a5", font=("Arial", 28))

    add_center_x = search_x - search_w - 45
    add_circle = canvas.create_oval(add_center_x - 23, 59, add_center_x + 23, 105, fill="#ffffff", outline="#d0d0d0", width=2)
    add_text = canvas.create_text(add_center_x, 82, text="+", fill="#000000", font=("Arial", 26, "bold"))

    def add_enter(event):
        canvas.itemconfig(add_circle, fill="#fafafa")
        root.config(cursor="hand2")

    def add_leave(event):
        canvas.itemconfig(add_circle, fill="#ffffff")
        root.config(cursor="")

    def add_click(event):
        show_dynamic_card_editor()

    for item in (add_circle, add_text):
        canvas.tag_bind(item, "<Enter>", add_enter)
        canvas.tag_bind(item, "<Leave>", add_leave)
        canvas.tag_bind(item, "<Button-1>", add_click)

    hidden = load_hidden_builtin_cards()
    dynamic_titles = [row["title"] for row in get_dynamic_cards()]
    all_items = [item for item in written_request_items if item not in hidden] + dynamic_titles

    query = written_request_search_query.strip().lower()
    if query:
        filtered_items = [item for item in all_items if query in item.lower()]
    else:
        filtered_items = list(all_items)

    visible_capacity = 30
    max_scroll = max(0, len(filtered_items) - visible_capacity)
    written_request_menu_scroll = max(0, min(written_request_menu_scroll, max_scroll))
    visible_items = filtered_items[written_request_menu_scroll:written_request_menu_scroll + visible_capacity]

    # Smaller cards: 3 columns x 10 rows
    columns = [
        visible_items[20:30],
        visible_items[10:20],
        visible_items[0:10],
    ]

    col_x = [310, 750, 1190]
    start_y = 145
    gap_y = 60
    card_w = 350
    card_h = 46

    for c, items in enumerate(columns):
        for r, title in enumerate(items):
            x = col_x[c]
            y = start_y + r * gap_y

            rounded_home_rect(x - card_w//2 + 6, y - card_h//2 + 8, x + card_w//2 + 6, y + card_h//2 + 8, r=8, fill="#cfcfcf", outline="#cfcfcf")

            card = rounded_home_rect(x - card_w//2, y - card_h//2, x + card_w//2, y + card_h//2, r=8, fill="#ffffff", outline="#ececec")

            size = 17
            sub = ""
            if title == "طلب توظيف عام":
                size = 16
                sub = "خارج إطار المسابقات"

            txt = canvas.create_text(x, y - 2 if sub else y, text=title, fill="#000000", font=("Arial", size, "bold"))

            subtxt = None
            if sub:
                subtxt = canvas.create_text(x, y + 14, text=sub, fill="#555555", font=("Arial", 9, "bold"))

            def enter(e, cd=card):
                canvas.itemconfig(cd, fill="#fafafa")
                root.config(cursor="hand2")

            def leave(e, cd=card):
                canvas.itemconfig(cd, fill="#ffffff")
                root.config(cursor="")

            def click(e, t=title):
                cancel_card_long_press()
                clear_dynamic_context_menu()
                job_form_entries["request_type"] = t
                if t == "طلب توظيف 1":
                    show_job_request_1_form()
                elif t == "طلب توظيف 2":
                    show_job_request_2_form()
                elif t == "مسابقة الجمارك":
                    show_customs_exam_form()
                elif t == "مسابقة الشرطة":
                    show_police_exam_form()
                elif t == "مسابقة الحماية المدنية":
                    show_civil_protection_exam_form()
                elif t == "عقود ماقبل التشغيل":
                    show_pre_employment_contract_form()
                elif t == "مسابقة الماستر":
                    show_master_exam_form()
                else:
                    dynamic_card_rows = db_fetchall("SELECT id FROM dynamic_cards WHERE title=? AND is_visible=1", (t,))
                    if dynamic_card_rows:
                        show_dynamic_card_form(dynamic_card_rows[0]["id"])
                    else:
                        show_job_request_form()

            def right_click(e, t=title):
                cancel_card_long_press()
                show_card_context_menu(e.x, e.y, t)

            def press(e, t=title):
                start_card_long_press(e, t)

            bind_items = [card, txt]
            if subtxt:
                bind_items.append(subtxt)

            for it in bind_items:
                canvas.tag_bind(it, "<Enter>", enter)
                canvas.tag_bind(it, "<Leave>", leave)
                canvas.tag_bind(it, "<Button-1>", click)
                canvas.tag_bind(it, "<Button-3>", right_click)
                canvas.tag_bind(it, "<ButtonPress-1>", press)
                canvas.tag_bind(it, "<ButtonRelease-1>", cancel_card_long_press)

    if not visible_items:
        canvas.create_text(760, 395, text="لا توجد نتائج", fill="#777777", font=("Arial", 28, "bold"))

    if written_request_menu_scroll > 0:
        canvas.create_text(width - 55, 145, text="▲", fill="#777777", font=("Arial", 20, "bold"))
    if written_request_menu_scroll < max_scroll:
        canvas.create_text(width - 55, height - 45, text="▼", fill="#777777", font=("Arial", 20, "bold"))


def scroll_written_request_menu(event):
    global written_request_menu_scroll
    if current_page != "written_request_menu":
        return

    hidden = load_hidden_builtin_cards()
    dynamic_titles = [row["title"] for row in get_dynamic_cards()]
    all_items = [item for item in written_request_items if item not in hidden] + dynamic_titles

    query = written_request_search_query.strip().lower()
    if query:
        filtered_items = [item for item in all_items if query in item.lower()]
    else:
        filtered_items = list(all_items)

    max_scroll = max(0, len(filtered_items) - 30)

    if event.delta < 0:
        written_request_menu_scroll += 3
    else:
        written_request_menu_scroll -= 3

    written_request_menu_scroll = max(0, min(written_request_menu_scroll, max_scroll))
    show_written_request_menu()


root.bind("<Configure>", on_resize)
root.bind("<Up>", on_request_type_key)
root.bind("<Down>", on_request_type_key)
root.bind("<Return>", on_request_type_key)
root.bind("<Escape>", on_request_type_key)
root.bind("<MouseWheel>", lambda event: scroll_written_request_menu(event) if current_page == "written_request_menu" else (scroll_request_type_dropdown(1 if event.delta < 0 else -1) if (current_page == "job_request_form" and request_type_dropdown_open) else (scroll_job_form(event) if current_page == "job_request_form" else scroll_written_request(event))))

show_home()
root.mainloop()
